import json
from io import BytesIO
from types import SimpleNamespace

import models
from docx import Document
from PIL import Image, ImageDraw
from services.evaluation_service import COMMON_DIMENSIONS
from services.workflow_service import workflow_service


CASE_TEXT = (
    "2026年5月1日21时许，报警人李娟称在XX路东段废弃仓库发现一名男子倒地，胸口有血迹。"
    "民警到场后确认男子已死亡。经调查，嫌疑人张磊与被害人王浩因债务纠纷发生冲突。"
)


class TestCasesAuth:
    def test_parse_requires_login(self, client):
        response = client.post("/cases/parse", json={"text": CASE_TEXT, "source_mode": "plain_case"})
        assert response.status_code == 401

    def test_parse_rejects_student(self, client, student_headers):
        response = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=student_headers,
        )
        assert response.status_code == 403

    def test_read_all_roles_rejects_student(self, client, student_headers):
        response = client.get("/cases/all/roles", headers=student_headers)
        assert response.status_code == 403


class TestCasesParse:
    def test_parse_plain_case(self, client, admin_headers):
        response = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=admin_headers,
        )
        assert response.status_code == 200

        data = response.json()
        assert data["case_type"] == "故意杀人"
        assert data["source_mode"] == "plain_case"
        assert data["parse_engine"] in {"ai", "heuristic"}
        assert data["fact_sheet"]["case_time"] == "2026年5月1日21时许"
        assert data["fact_sheet"]["case_location"] == "XX路东段废弃仓库"
        assert "XX路东段废弃仓库" in data["dispatch_brief_suggestion"]

        person_names = {item["name"] for item in data["persons"]}
        assert {"李娟", "张磊", "王浩"}.issubset(person_names)
        suspect = next(item for item in data["persons"] if item["name"] == "张磊")
        assert suspect["self_image"]
        assert suspect["current_need"]
        assert suspect["public_mask"]
        assert suspect["private_drive"]

    def test_normalize_parsed_case_cleans_person_name_suffixes(self):
        payload = {
            "persons": [
                {"name": "报警人李娟称", "role": "报警人"},
                {"name": "嫌疑人张磊因债务纠纷", "role": "嫌疑人"},
                {"name": "被害人王浩表示", "role": "被害人"},
            ]
        }

        result = workflow_service._normalize_parsed_case(CASE_TEXT, payload, "plain_case", None)
        person_names = {item["name"] for item in result["persons"]}
        assert {"李娟", "张磊", "王浩"}.issubset(person_names)
        assert "报警人李娟称" not in person_names
        assert "嫌疑人张磊因债务纠纷" not in person_names

    def test_parse_ai_result_with_unknown_numeric_person_scores_stays_ai(self, monkeypatch):
        def fake_completion(**kwargs):
            payload = {
                "case_name": "仓库发现尸体警情",
                "case_type": "故意杀人",
                "case_background": "报警人李娟在废弃仓库发现王浩倒地，后续调查涉及张磊与王浩的债务纠纷。",
                "persons": [
                    {
                        "name": "李娟",
                        "role": "报警人",
                        "role_type": "证人",
                        "init_risk": "未明确",
                        "init_expression_clarity": "待核实",
                    }
                ],
                "parse_engine": "ai",
            }
            return SimpleNamespace(
                choices=[
                    SimpleNamespace(
                        message=SimpleNamespace(content=json.dumps(payload, ensure_ascii=False)),
                        finish_reason="stop",
                    )
                ]
            )

        monkeypatch.setattr("services.workflow_service.create_json_chat_completion", fake_completion)
        monkeypatch.setattr(workflow_service, "extract_case_person_names", lambda _text: ["李娟"])

        result = workflow_service.parse_case_text(CASE_TEXT, source_mode="plain_case")

        assert result["parse_engine"] == "ai"
        assert result["persons"][0]["name"] == "李娟"
        assert isinstance(result["persons"][0]["init_risk"], int)


class TestCasesRoles:
    def test_read_all_roles_returns_scene_links_and_case_meta(self, client, admin_headers):
        response = client.get("/cases/all/roles", headers=admin_headers)
        assert response.status_code == 200

        data = response.json()
        assert isinstance(data, list)
        target = next(item for item in data if item["name"] == "张某")
        assert target["case_title"] == "邻里纠纷测试案件"
        assert target["case_id"] is not None
        assert target["scene_ids"]
        assert target["primary_scene_id"] in target["scene_ids"]
        assert target["is_public"] is False

    def test_public_role_round_trip_preserves_compact_persona_fields(self, client, admin_headers):
        create_payload = {
            "name": "公共模板-李某",
            "role_type": "证人",
            "interaction_style": "观察型",
            "personality": "谨慎、怕牵连自己",
            "speaking_style": "先试探，再补充",
            "init_emotion": 42,
            "init_trust": 58,
            "status": "正常",
            "iq_level": "中等",
            "eq_level": "高",
            "lying_ability": "一般",
            "knows_facts": ["听到过争吵声"],
            "does_not_know": ["不知道谁先动手"],
            "hidden_truths": ["其实提前见过双方一次"],
            "behavior_archetype": "谨慎回避型",
            "police_attitude": "试探观望",
            "current_goal": "先别把自己卷进正式责任",
            "core_concern": "最怕单位知道后影响工作",
            "calming_points": ["先讲清不会乱定性", "先从自己确定看到的部分问起"],
            "relationship_pressure": ["护着儿子", "怕邻居继续报复"],
            "surface_stance": "我只说自己亲眼看到的部分",
            "pressure_response": "被逼急了会先回避，再一点点补细节",
            "trigger_points": ["儿子会不会被牵连", "单位知道这件事"],
        }

        create_response = client.post("/cases/roles", json=create_payload, headers=admin_headers)
        assert create_response.status_code == 200
        role_id = create_response.json()["id"]

        roles_response = client.get("/cases/all/roles", headers=admin_headers)
        assert roles_response.status_code == 200
        created_role = next(item for item in roles_response.json() if item["id"] == role_id)
        assert created_role["is_public"] is True
        assert created_role["behavior_archetype"] == create_payload["behavior_archetype"]
        assert created_role["police_attitude"] == create_payload["police_attitude"]
        assert created_role["current_goal"] == create_payload["current_goal"]
        assert created_role["core_concern"] == create_payload["core_concern"]
        assert created_role["calming_points"] == create_payload["calming_points"]
        assert created_role["relationship_pressure"] == create_payload["relationship_pressure"]
        assert created_role["surface_stance"] == create_payload["surface_stance"]
        assert created_role["pressure_response"] == create_payload["pressure_response"]
        assert created_role["trigger_points"] == create_payload["trigger_points"]

        update_response = client.put(
            f"/cases/roles/{role_id}",
            json={
                **create_payload,
                "name": "公共模板-李某-更新",
                "behavior_archetype": "防御切责型",
                "police_attitude": "防备排斥",
                "current_goal": "先把自己摘出去，再决定补充多少",
                "core_concern": "最怕孩子和工作一起受影响",
                "calming_points": ["先按时间线核实", "给其把话讲完整的台阶"],
                "relationship_pressure": ["护着孩子", "担心邻里继续闹大"],
                "surface_stance": "我可以配合，但只说确定部分",
                "pressure_response": "先观察警方掌握程度，再慢慢补细节",
                "trigger_points": ["孩子受牵连", "单位追责"],
            },
            headers=admin_headers,
        )
        assert update_response.status_code == 200

        updated_role = next(item for item in client.get("/cases/all/roles", headers=admin_headers).json() if item["id"] == role_id)
        assert updated_role["name"] == "公共模板-李某-更新"
        assert updated_role["behavior_archetype"] == "防御切责型"
        assert updated_role["police_attitude"] == "防备排斥"
        assert updated_role["current_goal"] == "先把自己摘出去，再决定补充多少"
        assert updated_role["core_concern"] == "最怕孩子和工作一起受影响"
        assert updated_role["calming_points"] == ["先按时间线核实", "给其把话讲完整的台阶"]
        assert updated_role["relationship_pressure"] == ["护着孩子", "担心邻里继续闹大"]
        assert updated_role["surface_stance"] == "我可以配合，但只说确定部分"
        assert updated_role["pressure_response"] == "先观察警方掌握程度，再慢慢补细节"
        assert updated_role["trigger_points"] == ["孩子受牵连", "单位追责"]
        assert updated_role["hidden_truths"] == ["其实提前见过双方一次"]

    def test_public_role_preserves_new_reaction_archetype(self, client, admin_headers):
        payload = {
            "name": "公共模板-沉默证人",
            "role_type": "证人",
            "interaction_style": "观察型",
            "personality": "害怕报复，不敢当众说明情况",
            "speaking_style": "短答、停顿多",
            "status": "正常",
            "behavior_archetype": "沉默恐惧型",
            "current_goal": "先确认自己不会被对方报复",
            "core_concern": "怕说出事实后被找麻烦",
            "trigger_points": ["要求当面对质", "公开记录其身份"],
            "calming_points": ["说明保护措施", "先问能确认的安全事实"],
            "pressure_response": "被逼指认时会沉默或改口",
            "surface_stance": "我没怎么看清，别让我当面说。",
        }

        create_response = client.post("/cases/roles", json=payload, headers=admin_headers)
        assert create_response.status_code == 200
        role_id = create_response.json()["id"]

        created_role = next(item for item in client.get("/cases/all/roles", headers=admin_headers).json() if item["id"] == role_id)
        assert created_role["behavior_archetype"] == "沉默恐惧型"
        assert created_role["current_goal"] == payload["current_goal"]
        assert created_role["trigger_points"] == payload["trigger_points"]
        assert created_role["calming_points"] == payload["calming_points"]
        assert created_role["pressure_response"] == payload["pressure_response"]
        assert created_role["surface_stance"] == payload["surface_stance"]

    def test_parse_markdown_file(self, client, admin_headers):
        markdown_content = (
            "# 仓库命案记录\n\n"
            "2026年5月1日21时许，报警人李娟称在XX路东段废弃仓库发现一名男子倒地。\n"
            "经调查，嫌疑人张磊与被害人王浩因债务纠纷发生冲突。"
        )

        response = client.post(
            "/cases/parse-file",
            files={"file": ("case.md", markdown_content.encode("utf-8"), "text/markdown")},
            data={"source_mode": "transcript_file"},
            headers=admin_headers,
        )
        assert response.status_code == 200

        data = response.json()
        assert data["source_mode"] == "transcript_file"
        assert data["parse_engine"] in {"ai", "heuristic"}
        assert data["source_file_name"] == "case.md"
        assert data["source_file_type"] == "MD"
        assert data["source_file_size"] > 0
        assert "仓库命案记录" in data["extracted_text_preview"]
        assert data["case_type"] == "故意杀人"

    def test_parse_file_falls_back_to_rules_when_ai_parser_fails(self, client, admin_headers, monkeypatch):
        def broken_ai_parse(*args, **kwargs):
            raise RuntimeError("mock ai parser failure")

        monkeypatch.setattr(workflow_service, "parse_case_text", broken_ai_parse)

        markdown_content = (
            "# 仓库命案记录\n\n"
            "2026年5月1日21时许，报警人李娟称在XX路东段废弃仓库发现一名男子倒地。\n"
            "经调查，嫌疑人张磊与被害人王浩因债务纠纷发生冲突，现场有血迹和监控。"
        )

        response = client.post(
            "/cases/parse-file",
            files={"file": ("case.md", markdown_content.encode("utf-8"), "text/markdown")},
            data={"source_mode": "transcript_file"},
            headers=admin_headers,
        )
        assert response.status_code == 200

        data = response.json()
        assert data["parse_engine"] == "heuristic"
        assert data["source_file_name"] == "case.md"
        assert data["ocr_method"] == "markdown_text_extract"
        assert data["ocr_engine"] == "markdown-normalizer"
        assert "mock ai parser failure" in "\n".join(data["parse_warnings"])
        assert "仓库命案记录" in data["extracted_text_full"]

    def test_parse_scanned_pdf_uses_page_ocr(self, client, admin_headers, monkeypatch):
        from services.document_extract_service import document_extract_service

        def fake_ocr_image_bytes(_image_bytes: bytes, *, image_name: str):
            assert image_name.startswith("pdf-page-")
            return "2026年5月1日21时许，报警人李娟称在XX路东段废弃仓库发现一名男子倒地。", []

        monkeypatch.setattr(document_extract_service, "_ocr_image_bytes", fake_ocr_image_bytes)

        image = Image.new("RGB", (500, 220), "white")
        draw = ImageDraw.Draw(image)
        draw.text((20, 80), "scanned case image", fill="black")
        buffer = BytesIO()
        image.save(buffer, format="PDF")

        response = client.post(
            "/cases/parse-file",
            files={"file": ("scan.pdf", buffer.getvalue(), "application/pdf")},
            data={"source_mode": "transcript_file"},
            headers=admin_headers,
        )
        assert response.status_code == 200

        data = response.json()
        assert data["ocr_method"] == "pdf_page_ocr"
        assert data["ocr_engine"] == "pypdfium2+paddleocr"
        assert data["parse_engine"] in {"ai", "heuristic"}
        assert "【PDF OCR识别结果】" in data["extracted_text_full"]

    def test_parse_file_includes_ocr_metadata(self, client, admin_headers):
        document = Document()
        document.add_paragraph("2026年5月1日21时许，报警人李娟称在XX路东段废弃仓库发现一名男子倒地。")
        document.add_paragraph("经调查，嫌疑人张磊与被害人王浩因债务纠纷发生冲突。")
        buffer = BytesIO()
        document.save(buffer)

        response = client.post(
            "/cases/parse-file",
            files={"file": ("case.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"source_mode": "transcript_file"},
            headers=admin_headers,
        )
        assert response.status_code == 200

        data = response.json()
        assert data["source_file_name"] == "case.docx"
        assert data["ocr_method"] == "docx_unified_ocr"
        assert data["ocr_engine"] == "python-docx+paddleocr"
        assert "extracted_text_full" in data
        assert "ocr_warnings" in data
        assert "【文档识别结果】" in data["extracted_text_full"]

    def test_update_case_syncs_structured_persons_into_roles(self, client, admin_headers, db_session):
        parsed = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=admin_headers,
        ).json()
        scenes = client.post(
            "/cases/generate-scenes",
            json={"case_info": parsed},
            headers=admin_headers,
        ).json()["scenes"]
        created_case = client.post(
            "/cases/full-create",
            json={"case": parsed, "scenes": scenes},
            headers=admin_headers,
        ).json()

        db_case = db_session.query(models.Case).filter(models.Case.id == created_case["id"]).first()
        target_scene = (
            db_session.query(models.Scene)
            .filter(models.Scene.case_id == db_case.id)
            .order_by(models.Scene.id.asc())
            .first()
        )
        assert target_scene is not None

        updated_structured = json.loads(db_case.structured_data or "{}")
        updated_structured["persons"] = [
            {
                "name": "张某",
                "role": "报警人",
                "role_type": "证人",
                "interaction_style": "情绪型",
                "personality": "急躁、较真、容易委屈上头",
                "speaking_style": "快、碎、带抱怨",
                "init_emotion": 85,
                "init_trust": 35,
                "status": "正常",
                "knows_facts": ["李某长期在楼道堆放杂物"],
                "does_not_know": ["李某为什么突然锁门不出"],
                "hidden_truths": ["自己在推搡时先伸了手"],
                "weakness": "一提到家人和物业就容易激动",
            },
            {
                "name": "李某",
                "role": "邻居",
                "role_type": "相关人员",
                "interaction_style": "对抗型",
                "personality": "嘴硬、爱争理、不愿吃亏",
                "speaking_style": "阴阳怪气、容易顶嘴",
                "init_emotion": 72,
                "init_trust": 20,
                "status": "正常",
                "knows_facts": ["自己确实在楼道放过杂物"],
                "does_not_know": [],
                "hidden_truths": ["和张某之前已经闹过几次不愉快"],
                "weakness": "怕物业和社区一起介入",
            },
        ]

        response = client.put(
            f"/cases/{db_case.id}",
            json={
                "case": {
                    "title": db_case.title,
                    "case_type": db_case.case_type,
                    "background": db_case.background,
                    "original_content": db_case.original_content,
                    "structured_data": updated_structured,
                },
                "scenes": [
                    {
                        "id": target_scene.id,
                        "name": target_scene.name,
                        "description": target_scene.description,
                        "difficulty": target_scene.difficulty,
                        "dispatch_brief": target_scene.dispatch_brief,
                        "first_impression": target_scene.first_impression,
                        "stages": json.loads(target_scene.stages or "[]"),
                    }
                ],
            },
            headers=admin_headers,
        )
        assert response.status_code == 200

        db_session.expire_all()
        updated_role = (
            db_session.query(models.Role)
            .filter(models.Role.case_id == db_case.id, models.Role.name == "张某")
            .first()
        )
        new_role = (
            db_session.query(models.Role)
            .filter(models.Role.case_id == db_case.id, models.Role.name == "李某")
            .first()
        )

        assert updated_role is not None
        assert updated_role.role_type == "证人"
        assert updated_role.interaction_style == "情绪型"
        assert updated_role.personality == "急躁、较真、容易委屈上头"

        assert new_role is not None
        assert new_role.role_type == "相关人员"
        assert new_role.interaction_style == "对抗型"

        linked_scene_ids = {
            row.scene_id
            for row in db_session.query(models.SceneRole).filter(models.SceneRole.role_id == new_role.id).all()
        }
        assert linked_scene_ids


class TestCasesSceneGeneration:
    def test_generate_scenes_from_parsed_case(self, client, admin_headers):
        parsed = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=admin_headers,
        ).json()

        response = client.post(
            "/cases/generate-scenes",
            json={"case_info": parsed},
            headers=admin_headers,
        )
        assert response.status_code == 200

        data = response.json()
        assert len(data["scenes"]) >= 2
        assert data["scene_generation_mode"].startswith(("ai", "fallback"))
        assert any("XX路东段废弃仓库" in scene["dispatch_brief"] for scene in data["scenes"])
        assert all(scene["roles"] for scene in data["scenes"])
        assert all(scene["stages"] for scene in data["scenes"])
        assert all("stage_goal" in scene["stages"][0] for scene in data["scenes"])


class TestCasesFullCreate:
    def test_parse_case_builds_richer_persona_fields(self, client, admin_headers):
        response = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=admin_headers,
        )
        assert response.status_code == 200
        parsed = response.json()
        persons = parsed.get("persons") or []
        assert persons
        sample_person = persons[0]
        assert sample_person.get("interaction_style")
        assert sample_person.get("personality")
        assert sample_person.get("speaking_style")
        assert sample_person.get("self_image")
        assert sample_person.get("current_need")
        assert sample_person.get("authority_attitude")
        assert sample_person.get("stress_response")
        assert isinstance(sample_person.get("trigger_topics"), list)
        assert isinstance(sample_person.get("coping_patterns"), list)

    def test_full_create_case_persists_case_scenes_and_roles(self, client, admin_headers, db_session):
        parsed = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=admin_headers,
        ).json()
        scenes = client.post(
            "/cases/generate-scenes",
            json={"case_info": parsed},
            headers=admin_headers,
        ).json()["scenes"]

        response = client.post(
            "/cases/full-create",
            json={"case": parsed, "scenes": scenes},
            headers=admin_headers,
        )
        assert response.status_code == 200

        created = response.json()
        assert created["case_type"] == "故意杀人"
        assert created["title"]

        db_case = db_session.query(models.Case).filter(models.Case.id == created["id"]).first()
        assert db_case is not None
        assert db_case.case_type == "故意杀人"
        assert db_case.original_content == CASE_TEXT

        structured = json.loads(db_case.structured_data or "{}")
        assert structured["source_mode"] == "plain_case"
        assert structured["original_content"] == CASE_TEXT
        assert structured["fact_sheet"]["case_location"] == "XX路东段废弃仓库"
        assert structured["persons"][0]["self_image"]
        assert structured["persons"][0]["current_need"]

        db_scenes = db_session.query(models.Scene).filter(models.Scene.case_id == db_case.id).all()
        db_roles = db_session.query(models.Role).filter(models.Role.case_id == db_case.id).all()
        scene_role_links = (
            db_session.query(models.SceneRole)
            .join(models.Scene, models.Scene.id == models.SceneRole.scene_id)
            .filter(models.Scene.case_id == db_case.id)
            .all()
        )

        assert len(db_scenes) >= 2
        assert len(db_roles) >= 3
        assert len(scene_role_links) >= 2
        assert any(link.is_primary for link in scene_role_links)
        assert structured.get("scene_role_map")

    def test_full_create_standardizes_person_names_and_scene_roles(self, client, admin_headers, db_session):
        response = client.post(
            "/cases/full-create",
            json={
                "case": {
                    "case_name": "姓名标准化测试",
                    "case_type": "其他",
                    "case_background": "报警人李四称张三在幸福小区与其发生争执。",
                    "original_content": "报警人李四称，张三在幸福小区与其发生争执。证人张三表示只看到争吵。",
                    "persons": [
                        {"name": "张三（审讯阶段）", "role_type": "嫌疑人", "status": "正常"},
                        {"name": "张三嫌疑人", "role_type": "嫌疑人", "status": "正常"},
                        {"name": "幸福小区", "role_type": "相关人员", "status": "正常"},
                        {"name": "口供", "role_type": "相关人员", "status": "正常"},
                    ],
                },
                "scenes": [
                    {
                        "scene_name": "重点询问",
                        "scene_description": "测试场景",
                        "difficulty": "中等",
                        "roles": ["张三嫌疑人", "幸福小区"],
                        "primary_role_name": "张三（审讯阶段）",
                        "stages": [{"stage_name": "核实身份", "stage_goal": "核实人员身份。"}],
                    }
                ],
            },
            headers=admin_headers,
        )

        assert response.status_code == 200
        created = response.json()
        db_case = db_session.query(models.Case).filter(models.Case.id == created["id"]).first()
        structured = json.loads(db_case.structured_data or "{}")
        assert [person["name"] for person in structured["persons"]] == ["张三"]
        assert structured["persons"][0]["person_id"] == "P001"
        assert structured["scene_role_map"]["重点询问"]["role_names"] == ["张三"]
        assert structured["scene_role_map"]["重点询问"]["primary_role_name"] == "张三"

        db_roles = db_session.query(models.Role).filter(models.Role.case_id == db_case.id).all()
        assert [role.name for role in db_roles] == ["张三"]


class TestCasesTrainingFlow:
    def test_admin_created_case_can_be_used_by_student_training_flow(self, client, admin_headers, db_session):
        parsed = client.post(
            "/cases/parse",
            json={"text": CASE_TEXT, "source_mode": "plain_case"},
            headers=admin_headers,
        ).json()
        scenes = client.post(
            "/cases/generate-scenes",
            json={"case_info": parsed},
            headers=admin_headers,
        ).json()["scenes"]
        created_case = client.post(
            "/cases/full-create",
            json={"case": parsed, "scenes": scenes},
            headers=admin_headers,
        ).json()

        db_case = db_session.query(models.Case).filter(models.Case.id == created_case["id"]).first()
        db_scene = (
            db_session.query(models.Scene)
            .filter(models.Scene.case_id == db_case.id)
            .order_by(models.Scene.id.asc())
            .first()
        )
        assert db_scene is not None

        student_login = client.post("/auth/token", data={"username": "student002", "password": "123456"})
        student_headers = {"Authorization": f"Bearer {student_login.json()['access_token']}"}

        start_response = client.post(f"/training/start/{db_scene.id}", headers=student_headers)
        assert start_response.status_code == 200
        session = start_response.json()
        assert session["scene_id"] == db_scene.id
        assert session["status"] == "active"

        session_detail = client.get(f"/training/session/{session['id']}", headers=student_headers)
        assert session_detail.status_code == 200
        session_data = session_detail.json()
        assert session_data["case_title"] == db_case.title
        assert session_data["scene_name"] == db_scene.name
        assert session_data["role_name"] in {"李娟", "张磊"}
        assert session_data["dispatch_brief"]

        chat_response = client.post(
            f"/training/chat/{session['id']}",
            json={"role": "user", "content": "请你先说明现场发生了什么，以及你和其他人的关系。"},
            headers=student_headers,
        )
        assert chat_response.status_code == 200
        chat_data = chat_response.json()
        assert chat_data["response"]
        assert isinstance(chat_data["recommended_questions"], list)

        finish_response = client.post(f"/training/finish/{session['id']}", headers=student_headers)
        assert finish_response.status_code == 200
        report = finish_response.json()
        assert "total_score" in report
        score_groups = {item.get("group") for item in report["scores"]}
        common_dimensions = {name for name, _ in COMMON_DIMENSIONS}
        report_dimensions = {item.get("dimension") for item in report["scores"]}
        assert common_dimensions.issubset(report_dimensions)
        assert "assessment" in score_groups
