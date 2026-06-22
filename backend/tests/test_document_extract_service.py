from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from services.document_extract_service import DocumentExtractService


def _docx_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_unified_recognition_preserves_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("询问笔录")
    document.add_paragraph("问：2026年5月1日你在哪里？")
    document.add_paragraph("答：我在XX路东段废弃仓库。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "报警人"
    table.cell(0, 1).text = "李娟"
    table.cell(1, 0).text = "嫌疑人"
    table.cell(1, 1).text = "张磊"

    result = DocumentExtractService().recognize_docx(_docx_bytes(document))

    assert result.method == "docx_unified_ocr"
    assert result.engine == "python-docx+paddleocr"
    assert "【文档识别结果】" in result.text
    assert "问：2026年5月1日你在哪里？" in result.text
    assert "答：我在XX路东段废弃仓库。" in result.text
    assert "[表格 body-1]" in result.text
    assert "报警人 | 李娟" in result.text
    assert result.metadata["native_text_chars"] > 0
    assert result.metadata["block_count"] >= 4


def test_docx_unified_recognition_reports_ocr_unavailable_for_images_but_keeps_text(monkeypatch):
    service = DocumentExtractService()

    monkeypatch.setattr(
        service,
        "_extract_docx_images",
        lambda _file_bytes: [("word/media/image1.png", b"fake image bytes")],
    )
    monkeypatch.setattr(
        service,
        "_ocr_image_bytes",
        lambda _image_bytes, image_name: ("", [f"OCR 引擎不可用，未识别图片 {image_name}：missing"]),
    )

    document = Document()
    document.add_paragraph("2026年5月1日21时许，报警人李娟称在XX路东段废弃仓库发现一名男子倒地。")

    result = service.recognize_docx(_docx_bytes(document))

    assert result.warnings
    assert any("OCR 引擎不可用" in warning for warning in result.warnings)
    assert "报警人李娟" in result.text


def test_docx_xml_text_fallback_keeps_table_text_when_body_blocks_empty(monkeypatch):
    service = DocumentExtractService()
    monkeypatch.setattr(service, "_iter_body_blocks", lambda _document: iter(()))
    monkeypatch.setattr(service, "_extract_docx_images", lambda _file_bytes: [])

    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "报警人"
    table.cell(0, 1).text = "李娟"
    table.cell(1, 0).text = "嫌疑人"
    table.cell(1, 1).text = "张磊"

    result = service.recognize_docx(_docx_bytes(document))

    assert result.method == "docx_unified_ocr"
    assert any("底层 XML 文本兜底" in warning for warning in result.warnings)
    assert "报警人" in result.text
    assert "李娟" in result.text
    assert "嫌疑人" in result.text
    assert "张磊" in result.text


def test_docx_xml_text_fallback_runs_when_table_iteration_fails(monkeypatch):
    service = DocumentExtractService()
    monkeypatch.setattr(service, "_extract_docx_images", lambda _file_bytes: [])

    def fail_table(*_args, **_kwargs):
        raise ValueError("在 tbtbl 中，顶层 tr is invalid")

    monkeypatch.setattr(service, "_format_table", fail_table)

    document = Document()
    document.add_paragraph("询问笔录")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "报警人"
    table.cell(0, 1).text = "李娟"
    table.cell(1, 0).text = "嫌疑人"
    table.cell(1, 1).text = "张磊"

    result = service.recognize_docx(_docx_bytes(document))

    assert result.method == "docx_unified_ocr"
    assert any("读取正文/表格失败" in warning for warning in result.warnings)
    assert any("追加 DOCX 底层 XML 文本" in warning for warning in result.warnings)
    assert "询问笔录" in result.text
    assert "报警人" in result.text
    assert "李娟" in result.text
    assert "嫌疑人" in result.text
    assert "张磊" in result.text


def test_empty_docx_without_text_or_images_is_rejected(monkeypatch):
    service = DocumentExtractService()
    monkeypatch.setattr(service, "_extract_docx_images", lambda _file_bytes: [])

    document = Document()

    with pytest.raises(ValueError, match="未识别到有效正文"):
        service.recognize_docx(_docx_bytes(document))
