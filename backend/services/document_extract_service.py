from __future__ import annotations

import hashlib
import os
import re
import tempfile
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class DocumentExtractionResult:
    text: str
    method: str
    engine: str
    warnings: list[str] = field(default_factory=list)
    blocks: list[dict[str, Any]] = field(default_factory=list)
    pages: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def preview(self) -> str:
        return self.text[:500]

    def as_source_meta(self, *, name: str, extension: str, size: int) -> dict[str, Any]:
        return {
            "name": name,
            "type": extension.lstrip(".").upper(),
            "size": size,
            "ocr_method": self.method,
            "ocr_engine": self.engine,
            "ocr_warnings": self.warnings,
            "ocr_blocks": len(self.blocks),
            "ocr_pages": len(self.pages),
            "ocr_metadata": self.metadata,
            "extracted_text_full": self.text,
            "extracted_text_preview": self.preview,
        }


class DocumentExtractService:
    MAX_FILE_SIZE = 20 * 1024 * 1024
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}
    MIN_MEANINGFUL_DOCX_TEXT_CHARS = 30
    OCR_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

    def __init__(self):
        self._ocr_engine: Any | None = None
        self._ocr_engine_error: Exception | None = None

    @staticmethod
    def _stable_id(prefix: str, value: str) -> str:
        digest = hashlib.sha1(value.encode("utf-8", errors="ignore")).hexdigest()[:10]
        return f"{prefix}-{digest}"

    @staticmethod
    def _block(kind: str, text: str, *, location: str, index: int, extra: dict[str, Any] | None = None) -> dict[str, Any]:
        return {
            "id": DocumentExtractService._stable_id(kind, f"{location}:{index}:{text[:80]}"),
            "type": kind,
            "location": location,
            "index": index,
            "text": str(text or "").strip(),
            **(extra or {}),
        }

    @staticmethod
    def _format_table(table, *, location: str, index: int) -> str:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = DocumentExtractService.normalize_inline_text(cell.text)
                cells.append(value)
            if any(cells):
                rows.append(" | ".join(cells))
        if not rows:
            return ""
        return f"[表格 {location}-{index}]\n" + "\n".join(rows)

    @staticmethod
    def _iter_body_blocks(document: DocxDocument):
        body = document.element.body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield "paragraph", Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield "table", Table(child, document)

    @staticmethod
    def normalize_inline_text(text: str) -> str:
        return re.sub(r"[ \t\u3000]+", " ", str(text or "").replace("\r", "\n")).strip()

    @staticmethod
    def _extract_docx_xml_text(file_bytes: bytes) -> list[str]:
        """Fallback text extraction for unusual DOCX table/body structures."""
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        parts: list[str] = []
        with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
            names = [
                name
                for name in archive.namelist()
                if name == "word/document.xml"
                or name.startswith("word/header")
                or name.startswith("word/footer")
                or name.startswith("word/footnotes")
                or name.startswith("word/endnotes")
            ]
            for name in names:
                try:
                    root = ET.fromstring(archive.read(name))
                except Exception:
                    continue
                for paragraph in root.findall(".//w:p", namespaces):
                    texts = [
                        item.text or ""
                        for item in paragraph.findall(".//w:t", namespaces)
                        if item.text
                    ]
                    line = DocumentExtractService.normalize_inline_text("".join(texts))
                    if line:
                        parts.append(line)
        return parts

    @staticmethod
    def extract_pdf_text(file_bytes: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError as exc:
            raise ValueError("当前环境未安装 PDF 解析依赖 pypdf，请先安装后再使用 PDF 导入") from exc

        reader = PdfReader(BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        text = "\n".join(texts).strip()
        if not text:
            raise ValueError("当前版本不支持图片 OCR，请上传可复制文字版 PDF 或 DOCX")
        return text

    def recognize_pdf(self, file_bytes: bytes) -> DocumentExtractionResult:
        try:
            raw_text = self.extract_pdf_text(file_bytes)
            text = self.normalize_extracted_text(raw_text)
            return DocumentExtractionResult(
                text=text,
                method="pdf_text_extract",
                engine="pypdf",
                blocks=[self._block("pdf_text", text, location="pdf", index=1)],
            )
        except ValueError as text_exc:
            warnings = [f"PDF 可复制文本提取失败，已尝试 OCR：{text_exc}"]

        try:
            import pypdfium2 as pdfium
        except ModuleNotFoundError as exc:
            raise ValueError("PDF 未提取到可复制文字，且当前环境缺少 pypdfium2，无法对扫描件做 OCR") from exc

        blocks: list[dict[str, Any]] = []
        page_count = 0
        try:
            pdf = pdfium.PdfDocument(file_bytes)
            page_count = len(pdf)
            for page_index in range(page_count):
                page = pdf[page_index]
                bitmap = page.render(scale=2).to_pil()
                image_buffer = BytesIO()
                bitmap.save(image_buffer, format="PNG")
                page_text, page_warnings = self._ocr_image_bytes(
                    image_buffer.getvalue(),
                    image_name=f"pdf-page-{page_index + 1}.png",
                )
                warnings.extend(page_warnings)
                if page_text:
                    blocks.append(
                        self._block(
                            "pdf_page_ocr",
                            f"[PDF页面OCR {page_index + 1}]\n{page_text}",
                            location=f"page-{page_index + 1}",
                            index=page_index + 1,
                            extra={"page_index": page_index + 1},
                        )
                    )
        except Exception as exc:
            raise ValueError(f"PDF 扫描件 OCR 失败：{exc}") from exc

        if not blocks:
            raise ValueError("PDF 未提取到可复制文字，扫描件 OCR 也未识别到有效文本，请检查文件是否清晰或改传 DOCX/TXT")

        formatted_parts = ["【PDF OCR识别结果】", "说明：该 PDF 未提取到可复制文字，以下内容由页面 OCR 识别生成，请人工复核。"]
        for index, block in enumerate(blocks, start=1):
            formatted_parts.append(f"\n--- 页面 {index} / {block['type']} / {block['location']} ---")
            formatted_parts.append(block["text"])

        text = self.normalize_extracted_text("\n".join(formatted_parts))
        return DocumentExtractionResult(
            text=text,
            method="pdf_page_ocr",
            engine="pypdfium2+paddleocr",
            warnings=warnings,
            blocks=blocks,
            pages=[{"page_index": index + 1} for index in range(page_count)],
            metadata={"page_count": page_count, "ocr_blocks": len(blocks)},
        )

    @staticmethod
    def extract_plain_text(file_bytes: bytes | str) -> str:
        if isinstance(file_bytes, str):
            text = file_bytes
            last_error = None
        else:
            last_error: Exception | None = None
            for encoding in ("utf-8-sig", "utf-8", "gb18030"):
                try:
                    text = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError as exc:
                    last_error = exc
            else:
                raise ValueError("TXT 文档编码无法识别，请使用 UTF-8 或 GB18030 编码") from last_error

        text = text.strip()
        if not text:
            raise ValueError("TXT 文档未提取到有效正文，请检查文件内容后重试")
        return text

    @staticmethod
    def extract_markdown_text(file_bytes: bytes | str) -> str:
        if isinstance(file_bytes, str):
            text = file_bytes
            last_error = None
        else:
            last_error: Exception | None = None
            for encoding in ("utf-8-sig", "utf-8", "gb18030"):
                try:
                    text = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError as exc:
                    last_error = exc
            else:
                raise ValueError("Markdown 文档编码无法识别，请使用 UTF-8 或 GB18030 编码") from last_error

        text = re.sub(r"^\s{0,3}(#{1,6}|\- |\* |\+ |\d+\.)\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"`{1,3}", "", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = text.strip()
        if not text:
            raise ValueError("Markdown 文档未提取到有效正文，请检查文件内容后重试")
        return text

    @staticmethod
    def normalize_extracted_text(text: str) -> str:
        cleaned = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned)
        cleaned = re.sub(r"^\s*\d+\s*$", "", cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r"^\s*第\s*\d+\s*页\s*$", "", cleaned, flags=re.MULTILINE)
        return cleaned.strip()

    def _load_paddle_ocr(self):
        if self._ocr_engine is not None:
            return self._ocr_engine
        if self._ocr_engine_error is not None:
            raise self._ocr_engine_error
        try:
            from paddleocr import PaddleOCR

            try:
                self._ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
            except ValueError as exc:
                if "show_log" not in str(exc):
                    raise
                self._ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch")
            return self._ocr_engine
        except Exception as exc:
            self._ocr_engine_error = exc
            raise

    def _ocr_image_bytes(self, image_bytes: bytes, *, image_name: str) -> tuple[str, list[str]]:
        warnings: list[str] = []
        if not image_bytes:
            return "", ["图片内容为空，已跳过 OCR。"]

        suffix = Path(image_name).suffix.lower() or ".png"
        try:
            ocr = self._load_paddle_ocr()
        except Exception as exc:
            return "", [f"OCR 引擎不可用，未识别图片 {image_name}：{exc}"]

        temp_path = ""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                temp_file.write(image_bytes)
                temp_path = temp_file.name
            try:
                raw_result = ocr.ocr(temp_path, cls=True)
            except TypeError as exc:
                if "cls" not in str(exc):
                    raise
                raw_result = ocr.ocr(temp_path)
            lines: list[str] = []
            low_confidence_count = 0
            for page in raw_result or []:
                for item in page or []:
                    if not item or len(item) < 2:
                        continue
                    text_score = item[1]
                    if not isinstance(text_score, (list, tuple)) or not text_score:
                        continue
                    text = str(text_score[0] or "").strip()
                    score = float(text_score[1]) if len(text_score) > 1 else 0.0
                    if text:
                        lines.append(text)
                    if score and score < 0.65:
                        low_confidence_count += 1
            if low_confidence_count:
                warnings.append(f"图片 {image_name} 存在 {low_confidence_count} 条低置信 OCR 结果，请人工复核。")
            return "\n".join(lines).strip(), warnings
        except Exception as exc:
            return "", [f"OCR 识别图片 {image_name} 失败：{exc}"]
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except OSError:
                    pass

    def _extract_docx_images(self, file_bytes: bytes) -> list[tuple[str, bytes]]:
        images: list[tuple[str, bytes]] = []
        with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
            for name in archive.namelist():
                if not name.startswith("word/media/"):
                    continue
                suffix = Path(name).suffix.lower()
                if suffix not in self.OCR_IMAGE_EXTENSIONS:
                    continue
                images.append((name, archive.read(name)))
        return images

    def _append_docx_xml_blocks(
        self,
        file_bytes: bytes,
        blocks: list[dict[str, Any]],
        *,
        start_index: int,
    ) -> int:
        block_index = start_index
        for line in self._extract_docx_xml_text(file_bytes):
            blocks.append(
                self._block(
                    "docx_xml_text",
                    line,
                    location="docx_xml",
                    index=block_index,
                )
            )
            block_index += 1
        return block_index

    def recognize_docx(self, file_bytes: bytes) -> DocumentExtractionResult:
        warnings: list[str] = []
        try:
            document = Document(BytesIO(file_bytes))
        except Exception as exc:
            xml_lines = self._extract_docx_xml_text(file_bytes)
            if not xml_lines:
                raise ValueError(f"Word 文档结构读取失败，且未从底层 XML 提取到有效文字：{exc}") from exc
            text = self.normalize_extracted_text("\n".join(["【DOCX底层文本识别结果】", *xml_lines]))
            return DocumentExtractionResult(
                text=text,
                method="docx_xml_text_fallback",
                engine="docx-xml",
                warnings=[f"python-docx 读取失败，已使用 DOCX 底层 XML 文本兜底：{exc}"],
                blocks=[
                    self._block("docx_xml_text", line, location="docx_xml", index=index)
                    for index, line in enumerate(xml_lines, start=1)
                ],
                metadata={"xml_text_lines": len(xml_lines), "native_text_chars": len(text)},
            )

        blocks: list[dict[str, Any]] = []
        block_index = 1

        try:
            for section_index, section in enumerate(document.sections, start=1):
                for paragraph in section.header.paragraphs:
                    text = self.normalize_inline_text(paragraph.text)
                    if text:
                        blocks.append(self._block("header", f"[页眉] {text}", location=f"section-{section_index}", index=block_index))
                        block_index += 1
                for paragraph in section.footer.paragraphs:
                    text = self.normalize_inline_text(paragraph.text)
                    if text:
                        blocks.append(self._block("footer", f"[页脚] {text}", location=f"section-{section_index}", index=block_index))
                        block_index += 1
        except Exception as exc:
            warnings.append(f"python-docx 读取页眉页脚失败，已跳过该部分：{exc}")

        paragraph_index = 0
        table_index = 0
        body_read_failed = False
        try:
            for block_type, block in self._iter_body_blocks(document):
                if block_type == "paragraph":
                    paragraph_index += 1
                    text = self.normalize_inline_text(block.text)
                    if not text:
                        continue
                    style_name = str(getattr(block.style, "name", "") or "").strip()
                    prefix = f"[{style_name}] " if style_name and style_name.lower().startswith("heading") else ""
                    blocks.append(
                        self._block(
                            "paragraph",
                            f"{prefix}{text}",
                            location="body",
                            index=block_index,
                            extra={"paragraph_index": paragraph_index, "style": style_name},
                        )
                    )
                    block_index += 1
                    continue

                table_index += 1
                table_text = self._format_table(block, location="body", index=table_index)
                if table_text:
                    blocks.append(
                        self._block(
                            "table",
                            table_text,
                            location="body",
                            index=block_index,
                            extra={"table_index": table_index},
                        )
                    )
                    block_index += 1
        except Exception as exc:
            body_read_failed = True
            warnings.append(f"python-docx 读取正文/表格失败，已使用 DOCX 底层 XML 文本兜底：{exc}")
            before_count = len(blocks)
            block_index = self._append_docx_xml_blocks(file_bytes, blocks, start_index=block_index)
            if len(blocks) == before_count:
                warnings.append("DOCX 底层 XML 文本兜底未提取到额外正文。")

        try:
            docx_images = self._extract_docx_images(file_bytes)
        except Exception as exc:
            docx_images = []
            warnings.append(f"DOCX 图片抽取失败，已跳过图片 OCR：{exc}")

        for image_index, (image_name, image_bytes) in enumerate(docx_images, start=1):
            image_text, image_warnings = self._ocr_image_bytes(image_bytes, image_name=image_name)
            warnings.extend(image_warnings)
            if image_text:
                blocks.append(
                    self._block(
                        "image_ocr",
                        f"[图片OCR {image_index}: {image_name}]\n{image_text}",
                        location="body",
                        index=block_index,
                        extra={"image_index": image_index, "image_name": image_name},
                    )
                )
                block_index += 1

        if not blocks:
            block_index = self._append_docx_xml_blocks(file_bytes, blocks, start_index=block_index)
            if blocks:
                warnings.append("python-docx 未读取到正文块，已使用 DOCX 底层 XML 文本兜底。")
        elif body_read_failed:
            warnings.append("已保留 python-docx 先前读取到的内容，并追加 DOCX 底层 XML 文本供 AI/规则解析。")

        if not blocks:
            raise ValueError("Word 文档未识别到有效正文或图片文字，请检查文件内容后重试")

        native_chars = sum(len(block["text"]) for block in blocks if block.get("type") != "image_ocr")
        image_ocr_blocks = [block for block in blocks if block.get("type") == "image_ocr"]
        if native_chars < self.MIN_MEANINGFUL_DOCX_TEXT_CHARS and not image_ocr_blocks:
            warnings.append("DOCX 可复制正文较少，且未取得图片 OCR 文本，识别结果可能不完整。")

        formatted_parts = ["【文档识别结果】", "说明：以下内容按 DOCX 可识别顺序整理，保留段落、表格、页眉页脚和图片 OCR 标记；未进行摘要或改写。"]
        for index, block in enumerate(blocks, start=1):
            formatted_parts.append(f"\n--- 块 {index} / {block['type']} / {block['location']} ---")
            formatted_parts.append(block["text"])

        text = self.normalize_extracted_text("\n".join(formatted_parts))
        return DocumentExtractionResult(
            text=text,
            method="docx_unified_ocr",
            engine="python-docx+paddleocr",
            warnings=warnings,
            blocks=blocks,
            pages=[],
            metadata={
                "native_text_chars": native_chars,
                "image_ocr_blocks": len(image_ocr_blocks),
                "block_count": len(blocks),
            },
        )

    def recognize_file(self, filename: str, file_bytes: bytes) -> DocumentExtractionResult:
        lower_name = (filename or "").lower()
        if lower_name.endswith(".docx"):
            return self.recognize_docx(file_bytes)
        if lower_name.endswith(".pdf"):
            return self.recognize_pdf(file_bytes)
        if lower_name.endswith(".txt"):
            raw_text = self.extract_plain_text(file_bytes)
            text = self.normalize_extracted_text(raw_text)
            return DocumentExtractionResult(
                text=text,
                method="plain_text_extract",
                engine="text-decoder",
                blocks=[self._block("plain_text", text, location="txt", index=1)],
            )
        if lower_name.endswith(".md") or lower_name.endswith(".markdown"):
            raw_text = self.extract_markdown_text(file_bytes)
            text = self.normalize_extracted_text(raw_text)
            return DocumentExtractionResult(
                text=text,
                method="markdown_text_extract",
                engine="markdown-normalizer",
                blocks=[self._block("markdown_text", text, location="markdown", index=1)],
            )
        raise ValueError("仅支持 PDF、DOCX、TXT、MD 文件")

    def extract_text(self, filename: str, file_bytes: bytes) -> str:
        result = self.recognize_file(filename, file_bytes)
        if not result.text:
            raise ValueError("文档未提取到有效正文，请检查文件内容后重试")
        return result.text


document_extract_service = DocumentExtractService()
