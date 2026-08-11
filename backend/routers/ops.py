import csv
import io
import json
import os
import re
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy import bindparam, func, inspect, select, text
from sqlalchemy.orm import Session

import database
import models
import schemas
from routers.auth import build_username_range, hash_password, require_maintainer_user
from services.object_storage_service import delete_media_assets
from services.tabular_import_service import ExcelImportError, parse_excel_table

router = APIRouter(prefix="/ops", tags=["Ops"])

MANAGED_ROLES = {"admin", "student"}
VISIBLE_ROLES = {"admin", "student", "maintainer"}
FACE_TERMINATION_FAILURE_LIMIT = int(os.getenv("FACE_CONSECUTIVE_MAX_FAILURES", os.getenv("FACE_MAX_FAILURES", "5")))
IMPORT_FIELDS = ("username", "password", "role", "display_name", "real_name", "phone", "email", "unit", "department", "account_group", "bio")
IMPORT_HEADER_ALIASES = {
    "username": {"username", "user", "account", "账号", "账户", "用户名", "登录名", "学号", "工号"},
    "password": {"password", "pass", "pwd", "密码", "初始密码"},
    "role": {"role", "角色", "账号类型", "账户类型", "类型"},
    "display_name": {"display_name", "displayname", "name", "显示名", "昵称"},
    "real_name": {"real_name", "realname", "姓名", "真实姓名"},
    "phone": {"phone", "mobile", "tel", "手机号", "手机", "电话", "联系方式"},
    "email": {"email", "mail", "邮箱", "电子邮箱"},
    "unit": {"unit", "organization", "org", "单位", "机构", "学校"},
    "department": {"department", "dept", "部门", "班级", "院系"},
    "account_group": {"account_group", "group", "category", "分类", "分组", "账号分类", "账号分组", "开通批次", "批次"},
    "bio": {"bio", "remark", "remarks", "note", "notes", "备注", "说明"},
}
ROLE_ALIASES = {
    "admin": "admin",
    "student": "student",
    "管理端账号": "admin",
    "管理员": "admin",
    "管理": "admin",
    "学员端账号": "student",
    "学员": "student",
    "学生": "student",
}

PROFILE_FIELD_LIMITS = {
    "display_name": 80,
    "real_name": 80,
    "phone": 30,
    "email": 120,
    "unit": 120,
    "department": 120,
    "account_group": 80,
    "bio": 300,
}


def clean_text(value: str | None, max_length: int) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None
    if len(text) > max_length:
        raise HTTPException(status_code=400, detail=f"字段长度不能超过 {max_length} 个字符")
    return text


def cell_text(value) -> str:
    text = str(value if value is not None else "").strip()
    return re.sub(r"\s+", " ", text)


def normalize_header(value: str) -> str | None:
    raw = cell_text(value).strip().lower().replace(" ", "").replace("-", "_")
    if not raw:
        return None
    for field, aliases in IMPORT_HEADER_ALIASES.items():
        normalized_aliases = {alias.lower().replace(" ", "").replace("-", "_") for alias in aliases}
        if raw in normalized_aliases:
            return field
    return None


def table_to_records(rows: list[list[str]]) -> list[dict]:
    clean_rows = [[cell_text(cell) for cell in row] for row in rows if any(cell_text(cell) for cell in row)]
    if not clean_rows:
        return []

    header_index = -1
    headers: list[str | None] = []
    for index, row in enumerate(clean_rows[:10]):
        candidate = [normalize_header(cell) for cell in row]
        if "username" in candidate:
            header_index = index
            headers = candidate
            break
    if header_index < 0:
        raise HTTPException(status_code=400, detail="未找到账号导入表头，请至少提供 username/账号/学号 列")

    records: list[dict] = []
    for offset, row in enumerate(clean_rows[header_index + 1 :], start=header_index + 2):
        record = {"row_number": offset}
        for col_index, field in enumerate(headers):
            if field and field in IMPORT_FIELDS:
                record[field] = cell_text(row[col_index]) if col_index < len(row) else ""
        if any(record.get(field) for field in IMPORT_FIELDS):
            records.append(record)
    return records


def parse_csv_like(content: bytes, delimiter: str | None = None) -> list[dict]:
    last_error = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = content.decode(encoding)
            sample = text[:2048]
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;") if delimiter is None else csv.excel()
            if delimiter is not None:
                dialect.delimiter = delimiter
            return table_to_records([row for row in csv.reader(io.StringIO(text), dialect)])
        except Exception as exc:
            last_error = exc
    raise HTTPException(status_code=400, detail=f"无法解析文本表格文件：{last_error}")


def parse_xlsx(content: bytes) -> list[dict]:
    try:
        parsed = parse_excel_table(
            content,
            aliases=IMPORT_HEADER_ALIASES,
            required_field="username",
            allowed_fields=IMPORT_FIELDS,
        )
        return [
            {key: value for key, value in record.items() if key != "sheet_name"}
            for record in parsed.records
        ]
    except ExcelImportError as error:
        raise HTTPException(status_code=400, detail=str(error))


def parse_ods(content: bytes) -> list[dict]:
    namespaces = {"table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0", "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            root = ET.fromstring(archive.read("content.xml"))
            rows: list[list[str]] = []
            for row in root.findall(".//table:table-row", namespaces):
                values: list[str] = []
                for cell in row.findall("table:table-cell", namespaces):
                    repeat = int(cell.attrib.get("{urn:oasis:names:tc:opendocument:xmlns:table:1.0}number-columns-repeated", "1"))
                    text = " ".join(node.text or "" for node in cell.findall(".//text:p", namespaces))
                    values.extend([text] * min(repeat, 20))
                rows.append(values)
            return table_to_records(rows)
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="ODS 文件格式无效")


def parse_docx(content: bytes) -> list[dict]:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=400, detail="当前环境缺少 python-docx，无法解析 Word 文件")
    document = Document(io.BytesIO(content))
    tables = []
    for table in document.tables:
        tables.extend([[cell_text(cell.text) for cell in row.cells] for row in table.rows])
    if tables:
        return table_to_records(tables)
    lines = [line.split() for line in document.paragraphs if cell_text(line.text)]
    return table_to_records(lines)


def parse_pdf(content: bytes) -> list[dict]:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(status_code=400, detail="当前环境缺少 pypdf，无法解析 PDF 文件")
    reader = PdfReader(io.BytesIO(content))
    lines: list[list[str]] = []
    for page in reader.pages:
        for line in (page.extract_text() or "").splitlines():
            text = cell_text(line)
            if text:
                lines.append(re.split(r"\s{2,}|\t|,", text))
    return table_to_records(lines)


def parse_json_accounts(content: bytes) -> list[dict]:
    try:
        data = json.loads(content.decode("utf-8-sig"))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON 文件格式无效：{exc}")
    if isinstance(data, dict):
        data = data.get("accounts") or data.get("items") or []
    if not isinstance(data, list):
        raise HTTPException(status_code=400, detail="JSON 须为数组，或包含 accounts/items 数组")
    records = []
    for index, item in enumerate(data, start=1):
        if not isinstance(item, dict):
            continue
        record = {"row_number": index}
        for key, value in item.items():
            field = normalize_header(str(key)) or str(key)
            if field in IMPORT_FIELDS:
                record[field] = cell_text(value)
        records.append(record)
    return records


def parse_import_file(filename: str, content: bytes) -> list[dict]:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".csv":
        return parse_csv_like(content, ",")
    if suffix == ".tsv":
        return parse_csv_like(content, "\t")
    if suffix == ".txt":
        return parse_csv_like(content)
    if suffix == ".json":
        return parse_json_accounts(content)
    if suffix == ".xlsx":
        return parse_xlsx(content)
    if suffix == ".ods":
        return parse_ods(content)
    if suffix == ".docx":
        return parse_docx(content)
    if suffix == ".pdf":
        return parse_pdf(content)
    if suffix == ".xls":
        raise HTTPException(status_code=400, detail="暂不直接解析老式 .xls，请另存为 .xlsx 后导入")
    raise HTTPException(status_code=400, detail="不支持的文件格式，请使用 xlsx、csv、tsv、json、docx、ods 或文本型 pdf")


def serialize_account(user: models.User) -> schemas.OpsAccountOverview:
    return schemas.OpsAccountOverview.model_validate(user)


def delete_optional_table_rows(db: Session, table_name: str, column_name: str, values: list[int]) -> None:
    ids = [int(value) for value in values if int(value or 0) > 0]
    if not ids or not inspect(db.get_bind()).has_table(table_name):
        return
    statement = text(f'DELETE FROM "{table_name}" WHERE "{column_name}" IN :ids').bindparams(
        bindparam("ids", expanding=True)
    )
    db.execute(statement, {"ids": ids})


def delete_assignment_owned_data(db: Session, assignment_ids: list[int]) -> None:
    ids = [int(value) for value in assignment_ids if int(value or 0) > 0]
    if not ids:
        return
    db.query(models.AssignmentSubmission).filter(models.AssignmentSubmission.assignment_id.in_(ids)).delete(synchronize_session=False)
    db.query(models.AssignmentStudentOverride).filter(models.AssignmentStudentOverride.assignment_id.in_(ids)).delete(synchronize_session=False)
    db.query(models.TrainingAssignmentCase).filter(models.TrainingAssignmentCase.assignment_id.in_(ids)).delete(synchronize_session=False)
    db.query(models.TrainingAssignmentScene).filter(models.TrainingAssignmentScene.assignment_id.in_(ids)).delete(synchronize_session=False)
    db.query(models.TrainingAssignment).filter(models.TrainingAssignment.id.in_(ids)).delete(synchronize_session=False)


def delete_video_training_sessions(db: Session, session_ids: list[int]) -> None:
    ids = [int(value) for value in session_ids if int(value or 0) > 0]
    if not ids:
        return
    for session_id in ids:
        delete_media_assets(db, owner_type="training_session", owner_key=session_id)
    db.query(models.FaceVerificationEvent).filter(models.FaceVerificationEvent.video_session_id.in_(ids)).delete(synchronize_session=False)
    db.query(models.VideoNodeResult).filter(models.VideoNodeResult.session_id.in_(ids)).delete(synchronize_session=False)
    db.query(models.VideoTrainingSession).filter(models.VideoTrainingSession.id.in_(ids)).delete(synchronize_session=False)


def delete_video_owned_data(db: Session, video_ids: list[int]) -> None:
    ids = [int(value) for value in video_ids if int(value or 0) > 0]
    if not ids:
        return
    for video_id in ids:
        delete_media_assets(db, owner_type="video", owner_key=video_id)
    session_ids = [
        row[0]
        for row in db.query(models.VideoTrainingSession.id)
        .filter(models.VideoTrainingSession.video_id.in_(ids))
        .all()
    ]
    delete_video_training_sessions(db, session_ids)
    node_ids = [
        row[0]
        for row in db.query(models.VideoNode.id)
        .filter(models.VideoNode.video_id.in_(ids))
        .all()
    ]
    if node_ids:
        db.query(models.VideoNodeResult).filter(models.VideoNodeResult.node_id.in_(node_ids)).delete(synchronize_session=False)
        db.query(models.VideoNode).filter(models.VideoNode.id.in_(node_ids)).delete(synchronize_session=False)
    db.query(models.TrainingVideo).filter(models.TrainingVideo.id.in_(ids)).delete(synchronize_session=False)


def delete_admin_owned_data(db: Session, user_id: int) -> None:
    class_ids = [
        row[0]
        for row in db.query(models.TrainingClass.id)
        .filter(models.TrainingClass.created_by == user_id)
        .all()
    ]
    if class_ids:
        assignment_ids = [
            row[0]
            for row in db.query(models.TrainingAssignment.id)
            .filter(models.TrainingAssignment.class_id.in_(class_ids))
            .all()
        ]
        delete_assignment_owned_data(db, assignment_ids)
        db.query(models.ClassAnnouncement).filter(models.ClassAnnouncement.class_id.in_(class_ids)).delete(synchronize_session=False)
        db.query(models.ClassMembership).filter(models.ClassMembership.class_id.in_(class_ids)).delete(synchronize_session=False)
        db.query(models.TrainingClass).filter(models.TrainingClass.id.in_(class_ids)).delete(synchronize_session=False)

    assignment_ids = [
        row[0]
        for row in db.query(models.TrainingAssignment.id)
        .filter(models.TrainingAssignment.created_by == user_id)
        .all()
    ]
    delete_assignment_owned_data(db, assignment_ids)

    announcement_ids = [
        row[0]
        for row in db.query(models.ClassAnnouncement.id)
        .filter(models.ClassAnnouncement.created_by == user_id)
        .all()
    ]
    if announcement_ids:
        db.query(models.ClassAnnouncement).filter(models.ClassAnnouncement.id.in_(announcement_ids)).delete(synchronize_session=False)

    video_ids = [
        row[0]
        for row in db.query(models.TrainingVideo.id)
        .filter(models.TrainingVideo.uploaded_by == user_id)
        .all()
    ]
    delete_video_owned_data(db, video_ids)


def write_ops_audit(
    db: Session,
    *,
    actor: models.User | None,
    action: str,
    target_user: models.User | None = None,
    detail: dict | None = None,
) -> None:
    db.add(models.OpsAuditLog(
        actor_id=actor.id if actor else None,
        target_user_id=target_user.id if target_user else None,
        action=action,
        detail=json.dumps(detail or {}, ensure_ascii=False),
    ))


def ensure_account(db: Session, account_id: int) -> models.User:
    user = db.query(models.User).filter(models.User.id == account_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="账号不存在")
    return user


def apply_profile_fields(user: models.User, payload) -> None:
    for field, max_length in PROFILE_FIELD_LIMITS.items():
        if field in payload.model_fields_set:
            setattr(user, field, clean_text(getattr(payload, field), max_length))


def normalize_account_group_name(value: str | None) -> str:
    return clean_text(value, 80) or ""


def ensure_account_group_registry(db: Session, name: str | None, *, description: str | None = None) -> models.AccountGroup | None:
    group_name = normalize_account_group_name(name)
    if not group_name:
        return None
    group = db.query(models.AccountGroup).filter(models.AccountGroup.name == group_name).first()
    if group:
        if description is not None and clean_text(description, 200) is not None and group.description != clean_text(description, 200):
            group.description = clean_text(description, 200)
        return group
    group = models.AccountGroup(name=group_name, description=clean_text(description, 200))
    db.add(group)
    db.flush()
    return group


def normalize_import_role(role: str | None) -> str:
    raw = cell_text(role or "student")
    return ROLE_ALIASES.get(raw, ROLE_ALIASES.get(raw.lower(), raw.lower()))


def clean_import_record(record: dict) -> dict:
    cleaned = {"row_number": int(record.get("row_number") or 0)}
    for field in IMPORT_FIELDS:
        cleaned[field] = cell_text(record.get(field))
    cleaned["role"] = normalize_import_role(cleaned.get("role") or "student")
    if not cleaned["password"]:
        cleaned["password"] = "123456"
    return cleaned


def validate_import_records(records: list[dict], db: Session) -> list[schemas.OpsAccountImportPreviewItem]:
    existing_usernames = {
        row[0]
        for row in db.query(models.User.username).all()
    }
    seen: set[str] = set()
    items: list[schemas.OpsAccountImportPreviewItem] = []
    for record in records:
        item = clean_import_record(record)
        errors: list[str] = []
        username = item["username"]
        role = item["role"]
        if not username:
            errors.append("缺少账号")
        elif len(username) > 50:
            errors.append("账号不能超过 50 个字符")
        elif username in existing_usernames:
            errors.append("账号已存在")
        elif username in seen:
            errors.append("文件内账号重复")
        if username:
            seen.add(username)

        if len(item["password"]) < 6:
            errors.append("密码至少 6 位")
        if role not in MANAGED_ROLES:
            errors.append("角色只能是管理端账号或学员端账号")

        for field, max_length in (
            ("display_name", 80),
            ("real_name", 80),
            ("phone", 30),
            ("email", 120),
            ("unit", 120),
            ("department", 120),
            ("account_group", 80),
            ("bio", 300),
        ):
            if len(item.get(field) or "") > max_length:
                labels = {
                    "display_name": "显示名",
                    "real_name": "姓名",
                    "phone": "电话",
                    "email": "邮箱",
                    "unit": "单位",
                    "department": "部门",
                    "account_group": "分类",
                    "bio": "备注",
                }
                errors.append(f"{labels.get(field, field)}不能超过 {max_length} 个字符")

        items.append(schemas.OpsAccountImportPreviewItem(
            **item,
            status="error" if errors else "ready",
            errors=errors,
        ))
    return items


def normalize_target_role(value: str | None) -> str:
    role = str(value or "").strip()
    if not role:
        return "student"
    if role not in MANAGED_ROLES:
        raise HTTPException(status_code=400, detail="导入目标角色只能是管理端账号或学员端账号")
    return role


def apply_target_role(records: list[dict], target_role: str | None) -> list[dict]:
    role = normalize_target_role(target_role)
    cleaned: list[dict] = []
    for record in records:
        record = dict(record)
        if role in MANAGED_ROLES:
            record["role"] = role
        cleaned.append(record)
    return cleaned


def build_import_preview(filename: str, records: list[dict], db: Session) -> schemas.OpsAccountImportPreviewResponse:
    items = validate_import_records(records, db)
    ready_count = sum(1 for item in items if item.status == "ready")
    return schemas.OpsAccountImportPreviewResponse(
        filename=filename,
        total_count=len(items),
        ready_count=ready_count,
        error_count=len(items) - ready_count,
        items=items,
    )


def parse_json_detail(value: str | None) -> dict:
    if not value:
        return {}
    try:
        parsed = json.loads(value)
        return parsed if isinstance(parsed, dict) else {"value": parsed}
    except Exception:
        return {"raw": value}


def max_datetime(*values: datetime | None) -> datetime | None:
    clean = [value for value in values if value is not None]
    return max(clean) if clean else None


def timestamp_token(value: datetime | None) -> str:
    if not value:
        return ""
    return str(int(value.timestamp() * 1000))


def activity_item(category: str, module: str, action: str, title: str, created_at: datetime | None, detail: dict | None = None) -> dict:
    return {
        "category": category,
        "module": module,
        "action": action,
        "title": title,
        "created_at": created_at.isoformat() if created_at else None,
        "detail": detail or {},
    }


def face_usage_records(db: Session, user: models.User) -> list[dict]:
    rows: list[dict] = []

    passed_seen: set[int] = set()
    passed_events = (
        db.query(models.FaceVerificationEvent, models.TrainingSession, models.Scene, models.Case)
        .join(models.TrainingSession, models.TrainingSession.id == models.FaceVerificationEvent.session_id)
        .outerjoin(models.Scene, models.Scene.id == models.TrainingSession.scene_id)
        .outerjoin(models.Case, models.Case.id == models.Scene.case_id)
        .filter(
            models.FaceVerificationEvent.student_id == user.id,
            models.FaceVerificationEvent.event_type == "verify",
            models.FaceVerificationEvent.status == "passed",
        )
        .order_by(models.FaceVerificationEvent.created_at.asc(), models.FaceVerificationEvent.id.asc())
        .all()
    )
    for event, session, scene, case in passed_events:
        if event.session_id in passed_seen:
            continue
        passed_seen.add(event.session_id)
        scene_name = scene.name if scene else "未知场景"
        rows.append({
            "kind": "verify_passed",
            "event": event,
            "session": session,
            "scene": scene,
            "case": case,
            "title": f"人脸核验通过：{scene_name}",
            "description": "进入该案件场景时首次身份核验通过。",
            "status": "passed",
        })

    terminated_seen: set[int] = set()
    terminated_events = (
        db.query(models.FaceVerificationEvent, models.TrainingSession, models.Scene, models.Case)
        .join(models.TrainingSession, models.TrainingSession.id == models.FaceVerificationEvent.session_id)
        .outerjoin(models.Scene, models.Scene.id == models.TrainingSession.scene_id)
        .outerjoin(models.Case, models.Case.id == models.Scene.case_id)
        .filter(
            models.FaceVerificationEvent.student_id == user.id,
            models.FaceVerificationEvent.event_type != "verify",
            models.FaceVerificationEvent.status == "failed",
            models.FaceVerificationEvent.failure_count >= FACE_TERMINATION_FAILURE_LIMIT,
        )
        .order_by(models.FaceVerificationEvent.created_at.asc(), models.FaceVerificationEvent.id.asc())
        .all()
    )
    for event, session, scene, case in terminated_events:
        if event.session_id in terminated_seen:
            continue
        terminated_seen.add(event.session_id)
        scene_name = scene.name if scene else "未知场景"
        rows.append({
            "kind": "auto_terminated",
            "event": event,
            "session": session,
            "scene": scene,
            "case": case,
            "title": f"人脸异常自动结束：{scene_name}",
            "description": "训练中人脸异常累计达到上限，系统已结束本次训练。",
            "status": "terminated",
        })

    rows.sort(key=lambda item: item["event"].created_at or datetime.min, reverse=True)
    return rows


def latest_face_usage_record(db: Session, user: models.User) -> dict | None:
    records = face_usage_records(db, user)
    return records[0] if records else None


def account_usage_changed_at(db: Session, user: models.User) -> datetime | None:
    normal_session_ids = select(models.TrainingSession.id).filter(models.TrainingSession.user_id == user.id)
    video_session_ids = select(models.VideoTrainingSession.id).filter(models.VideoTrainingSession.user_id == user.id)
    return max_datetime(
        user.updated_at,
        user.last_login_at,
        user.created_at,
        db.query(func.max(models.OpsAuditLog.created_at)).filter(
            (models.OpsAuditLog.actor_id == user.id) | (models.OpsAuditLog.target_user_id == user.id)
        ).scalar(),
        db.query(func.max(models.TrainingClass.created_at)).filter(models.TrainingClass.created_by == user.id).scalar(),
        db.query(func.max(models.TrainingAssignment.created_at)).filter(models.TrainingAssignment.created_by == user.id).scalar(),
        db.query(func.max(models.TrainingAssignment.published_at)).filter(models.TrainingAssignment.created_by == user.id).scalar(),
        db.query(func.max(models.ClassAnnouncement.created_at)).filter(models.ClassAnnouncement.created_by == user.id).scalar(),
        db.query(func.max(models.TrainingVideo.created_at)).filter(models.TrainingVideo.uploaded_by == user.id).scalar(),
        db.query(func.max(models.TrainingVideo.updated_at)).filter(models.TrainingVideo.uploaded_by == user.id).scalar(),
        db.query(func.max(models.ClassMembership.joined_at)).filter(models.ClassMembership.user_id == user.id).scalar(),
        db.query(func.max(models.TrainingSession.created_at)).filter(models.TrainingSession.user_id == user.id).scalar(),
        db.query(func.max(models.TrainingSession.training_started_at)).filter(models.TrainingSession.user_id == user.id).scalar(),
        db.query(func.max(models.TrainingSession.training_finished_at)).filter(models.TrainingSession.user_id == user.id).scalar(),
        db.query(func.max(models.Message.created_at)).filter(models.Message.session_id.in_(normal_session_ids)).scalar(),
        db.query(func.max(models.TrainingSessionArtifact.created_at)).filter(models.TrainingSessionArtifact.session_id.in_(normal_session_ids)).scalar(),
        db.query(func.max(models.AssignmentSubmission.created_at)).filter(models.AssignmentSubmission.user_id == user.id).scalar(),
        db.query(func.max(models.AssignmentSubmission.updated_at)).filter(models.AssignmentSubmission.user_id == user.id).scalar(),
        db.query(func.max(models.AssignmentSubmission.submitted_at)).filter(models.AssignmentSubmission.user_id == user.id).scalar(),
        db.query(func.max(models.AssignmentStudentOverride.created_at)).filter(models.AssignmentStudentOverride.user_id == user.id).scalar(),
        db.query(func.max(models.AssignmentStudentOverride.updated_at)).filter(models.AssignmentStudentOverride.user_id == user.id).scalar(),
        db.query(func.max(models.VideoTrainingSession.created_at)).filter(models.VideoTrainingSession.user_id == user.id).scalar(),
        db.query(func.max(models.VideoTrainingSession.finished_at)).filter(models.VideoTrainingSession.user_id == user.id).scalar(),
        db.query(func.max(models.VideoTrainingSession.evaluation_completed_at)).filter(models.VideoTrainingSession.user_id == user.id).scalar(),
        db.query(func.max(models.VideoNodeResult.created_at)).filter(models.VideoNodeResult.session_id.in_(video_session_ids)).scalar(),
        db.query(func.max(models.FaceProfile.created_at)).filter(models.FaceProfile.student_id == user.id).scalar(),
        db.query(func.max(models.FaceProfile.updated_at)).filter(models.FaceProfile.student_id == user.id).scalar(),
        db.query(func.max(models.FaceVerificationEvent.created_at)).filter(models.FaceVerificationEvent.student_id == user.id).scalar(),
        db.query(func.max(models.SpeechUsageLog.created_at)).filter(models.SpeechUsageLog.user_id == user.id).scalar(),
    )


def serialize_usage_version(db: Session, user: models.User) -> dict:
    changed_at = account_usage_changed_at(db, user)
    return {
        "id": user.id,
        "version": timestamp_token(changed_at),
        "changed_at": changed_at.isoformat() if changed_at else None,
    }


def account_usage_stats(db: Session, user: models.User) -> dict:
    normal_sessions = db.query(models.TrainingSession).filter(models.TrainingSession.user_id == user.id)
    video_sessions = db.query(models.VideoTrainingSession).filter(models.VideoTrainingSession.user_id == user.id)
    submissions = db.query(models.AssignmentSubmission).filter(models.AssignmentSubmission.user_id == user.id)
    class_memberships = db.query(models.ClassMembership).filter(models.ClassMembership.user_id == user.id)

    admin_class_count = db.query(models.TrainingClass).filter(models.TrainingClass.created_by == user.id).count()
    admin_assignment_count = db.query(models.TrainingAssignment).filter(models.TrainingAssignment.created_by == user.id).count()
    admin_announcement_count = db.query(models.ClassAnnouncement).filter(models.ClassAnnouncement.created_by == user.id).count()
    admin_video_count = db.query(models.TrainingVideo).filter(models.TrainingVideo.uploaded_by == user.id).count()
    ops_action_count = db.query(models.OpsAuditLog).filter(
        (models.OpsAuditLog.actor_id == user.id) | (models.OpsAuditLog.target_user_id == user.id)
    ).count()
    message_count = (
        db.query(models.Message)
        .join(models.TrainingSession, models.TrainingSession.id == models.Message.session_id)
        .filter(models.TrainingSession.user_id == user.id)
        .count()
    )
    training_artifact_count = (
        db.query(models.TrainingSessionArtifact)
        .join(models.TrainingSession, models.TrainingSession.id == models.TrainingSessionArtifact.session_id)
        .filter(models.TrainingSession.user_id == user.id)
        .count()
    )
    face_event_count = len(face_usage_records(db, user))
    speech_usage_count = db.query(models.SpeechUsageLog).filter(models.SpeechUsageLog.user_id == user.id).count()
    speech_failed_count = db.query(models.SpeechUsageLog).filter(
        models.SpeechUsageLog.user_id == user.id,
        models.SpeechUsageLog.status == "failed",
    ).count()

    return {
        "normal_training_count": normal_sessions.count(),
        "normal_finished_count": normal_sessions.filter(models.TrainingSession.status == "finished").count(),
        "training_message_count": message_count,
        "training_artifact_count": training_artifact_count,
        "video_training_count": video_sessions.count(),
        "video_finished_count": video_sessions.filter(models.VideoTrainingSession.status == "finished").count(),
        "video_artifact_count": 0,
        "assignment_submission_count": submissions.count(),
        "class_join_count": class_memberships.count(),
        "admin_class_count": admin_class_count,
        "admin_assignment_count": admin_assignment_count,
        "admin_announcement_count": admin_announcement_count,
        "admin_video_count": admin_video_count,
        "ops_action_count": ops_action_count,
        "face_event_count": face_event_count,
        "speech_usage_count": speech_usage_count,
        "speech_failed_count": speech_failed_count,
    }


def account_profile(db: Session, user: models.User) -> dict:
    face_profile = db.query(models.FaceProfile).filter(models.FaceProfile.student_id == user.id).first()
    latest_face_record = latest_face_usage_record(db, user)
    latest_face_event = latest_face_record["event"] if latest_face_record else None
    class_rows = (
        db.query(models.ClassMembership, models.TrainingClass)
        .join(models.TrainingClass, models.TrainingClass.id == models.ClassMembership.class_id)
        .filter(models.ClassMembership.user_id == user.id)
        .order_by(models.ClassMembership.joined_at.desc())
        .limit(8)
        .all()
    )
    opened_users = (
        db.query(models.OpsAuditLog, models.User)
        .outerjoin(models.User, models.User.id == models.OpsAuditLog.target_user_id)
        .filter(
            models.OpsAuditLog.actor_id == user.id,
            models.OpsAuditLog.action.in_(("admin_register_account", "admin_batch_create_student", "admin_import_student", "create_account", "import_create_account")),
        )
        .order_by(models.OpsAuditLog.created_at.desc())
        .limit(20)
        .all()
    )
    samples = parse_json_detail(face_profile.sample_images_json).get("value") if face_profile and face_profile.sample_images_json else None
    if face_profile and not isinstance(samples, list):
        samples = [face_profile.face_image_url] if face_profile.face_image_url else []
    return {
        "id": user.id,
        "username": user.username,
        "role": user.role,
        "display_name": user.display_name,
        "real_name": user.real_name,
        "phone": user.phone,
        "email": user.email,
        "unit": user.unit,
        "department": user.department,
        "bio": user.bio,
        "created_at": user.created_at.isoformat() if user.created_at else None,
        "last_login_at": user.last_login_at.isoformat() if user.last_login_at else None,
        "face": {
            "registered": bool(face_profile),
            "created_at": face_profile.created_at.isoformat() if face_profile and face_profile.created_at else None,
            "updated_at": face_profile.updated_at.isoformat() if face_profile and face_profile.updated_at else None,
            "image_url": face_profile.face_image_url if face_profile else None,
            "sample_images": samples if isinstance(samples, list) else [],
            "sample_count": len(samples) if isinstance(samples, list) else None,
            "latest_status": latest_face_record["status"] if latest_face_record else None,
            "latest_reason": latest_face_record["description"] if latest_face_record else None,
            "latest_similarity": latest_face_event.similarity if latest_face_event else None,
            "latest_at": latest_face_event.created_at.isoformat() if latest_face_event and latest_face_event.created_at else None,
        },
        "classes": [
            {
                "class_id": classroom.id,
                "name": classroom.name,
                "role": membership.role,
                "status": membership.status,
                "joined_at": membership.joined_at.isoformat() if membership.joined_at else None,
            }
            for membership, classroom in class_rows
        ],
        "opened_accounts": [
            {
                "username": target.username if target else parse_json_detail(log.detail).get("username"),
                "role": target.role if target else parse_json_detail(log.detail).get("role"),
                "action": log.action,
                "created_at": log.created_at.isoformat() if log.created_at else None,
            }
            for log, target in opened_users
        ],
    }


def account_recent_activities(db: Session, user: models.User, limit: int = 30) -> list[dict]:
    activities: list[dict] = []
    if user.last_login_at:
        activities.append(activity_item("profile", "登录", "login", "账号登录系统", user.last_login_at, {"role": user.role}))

    for item in db.query(models.OpsAuditLog).filter(
        (models.OpsAuditLog.actor_id == user.id) | (models.OpsAuditLog.target_user_id == user.id)
    ).order_by(models.OpsAuditLog.created_at.desc()).limit(20).all():
        detail = parse_json_detail(item.detail)
        relation = "执行" if item.actor_id == user.id else "被操作"
        category = "admin" if item.action.startswith("admin_") else "audit"
        activities.append(activity_item(category, "账号审计", item.action, f"{relation}账号操作：{item.action}", item.created_at, detail))

    for item in db.query(models.TrainingClass).filter(models.TrainingClass.created_by == user.id).order_by(models.TrainingClass.created_at.desc()).limit(10):
        activities.append(activity_item("admin", "管理端", "create_class", f"创建班级：{item.name}", item.created_at, {"class_id": item.id}))

    for item in db.query(models.TrainingAssignment).filter(models.TrainingAssignment.created_by == user.id).order_by(models.TrainingAssignment.created_at.desc()).limit(10):
        activities.append(activity_item("admin", "管理端", "create_assignment", f"发布作业：{item.title}", item.created_at, {"assignment_id": item.id, "status": item.status}))

    for item in db.query(models.ClassAnnouncement).filter(models.ClassAnnouncement.created_by == user.id).order_by(models.ClassAnnouncement.created_at.desc()).limit(10):
        activities.append(activity_item("admin", "管理端", "create_announcement", f"发布公告：{item.title}", item.created_at, {"announcement_id": item.id}))

    for item in db.query(models.TrainingVideo).filter(models.TrainingVideo.uploaded_by == user.id).order_by(models.TrainingVideo.created_at.desc()).limit(10):
        activities.append(activity_item("admin", "管理端", "upload_video", f"上传视频：{item.title}", item.created_at, {"video_id": item.id, "status": item.status}))

    for membership, classroom in (
        db.query(models.ClassMembership, models.TrainingClass)
        .join(models.TrainingClass, models.TrainingClass.id == models.ClassMembership.class_id)
        .filter(models.ClassMembership.user_id == user.id)
        .order_by(models.ClassMembership.joined_at.desc())
        .limit(10)
        .all()
    ):
        activities.append(activity_item("assignment", "学员端", "join_class", f"加入班级：{classroom.name}", membership.joined_at, {
            "class_id": classroom.id,
            "membership_status": membership.status,
        }))

    face_profile = db.query(models.FaceProfile).filter(models.FaceProfile.student_id == user.id).first()
    if face_profile:
        activities.append(activity_item("face", "人脸", "face_profile", "人脸档案更新", face_profile.updated_at or face_profile.created_at, {
            "status": "registered",
            "image_url": face_profile.face_image_url,
            "sample_count": len(parse_json_detail(face_profile.sample_images_json).get("value") or []),
        }))

    normal_rows = (
        db.query(models.TrainingSession, models.Scene, models.Case)
        .outerjoin(models.Scene, models.Scene.id == models.TrainingSession.scene_id)
        .outerjoin(models.Case, models.Case.id == models.Scene.case_id)
        .filter(models.TrainingSession.user_id == user.id)
        .order_by(models.TrainingSession.created_at.desc())
        .limit(15)
        .all()
    )
    for session, scene, case in normal_rows:
        title = f"普通训练：{scene.name if scene else '未知场景'}"
        message_count = db.query(models.Message).filter(models.Message.session_id == session.id).count()
        artifact_count = db.query(models.TrainingSessionArtifact).filter(models.TrainingSessionArtifact.session_id == session.id).count()
        activities.append(activity_item("training", "训练", "training_session", title, session.created_at, {
            "session_id": session.id,
            "status": session.status,
            "case_title": case.title if case else None,
            "message_count": message_count,
            "artifact_count": artifact_count,
            "finished_at": session.training_finished_at.isoformat() if session.training_finished_at else None,
        }))

    video_rows = (
        db.query(models.VideoTrainingSession, models.TrainingVideo)
        .outerjoin(models.TrainingVideo, models.TrainingVideo.id == models.VideoTrainingSession.video_id)
        .filter(models.VideoTrainingSession.user_id == user.id)
        .order_by(models.VideoTrainingSession.created_at.desc())
        .limit(15)
        .all()
    )
    for session, video in video_rows:
        node_count = db.query(models.VideoNodeResult).filter(models.VideoNodeResult.session_id == session.id).count()
        activities.append(activity_item("video", "视频训练", "video_training_session", f"视频训练：{video.title if video else '未知视频'}", session.created_at, {
            "session_id": session.id,
            "status": session.status,
            "mode": session.mode,
            "score": session.total_score,
            "node_count": node_count,
            "finished_at": session.finished_at.isoformat() if session.finished_at else None,
        }))

    submission_rows = (
        db.query(models.AssignmentSubmission, models.TrainingAssignment)
        .outerjoin(models.TrainingAssignment, models.TrainingAssignment.id == models.AssignmentSubmission.assignment_id)
        .filter(models.AssignmentSubmission.user_id == user.id)
        .order_by(models.AssignmentSubmission.updated_at.desc())
        .limit(15)
        .all()
    )
    for submission, assignment in submission_rows:
        activities.append(activity_item("assignment", "作业", "assignment_submission", f"作业提交：{assignment.title if assignment else '未知作业'}", submission.updated_at, {
            "submission_id": submission.id,
            "assignment_id": submission.assignment_id,
            "case_id": submission.case_id,
            "scene_id": submission.scene_id,
            "status": submission.status,
            "score": submission.score,
            "submitted_at": submission.submitted_at.isoformat() if submission.submitted_at else None,
        }))

    for record in face_usage_records(db, user)[:20]:
        event = record["event"]
        session = record["session"]
        scene = record["scene"]
        case = record["case"]
        activities.append(activity_item("face", "人脸", "face_verification", record["title"], event.created_at, {
            "event_type": record["kind"],
            "status": record["status"],
            "reason": record["description"],
            "session_id": session.id if session else None,
            "scene_id": scene.id if scene else None,
            "case_id": case.id if case else None,
            "case_title": case.title if case else None,
            "similarity": event.similarity,
            "failure_count": event.failure_count if record["kind"] == "auto_terminated" else None,
        }))

    for item in db.query(models.SpeechUsageLog).filter(models.SpeechUsageLog.user_id == user.id).order_by(models.SpeechUsageLog.created_at.desc()).limit(30):
        mode_label = "实时语音识别" if item.mode == "realtime" else "语音文件识别"
        activities.append(activity_item("speech", "语音识别", "speech_usage", f"{mode_label}：{item.status}", item.created_at, {
            "mode": item.mode,
            "status": item.status,
            "language": item.language,
            "model": item.model,
            "duration_seconds": item.duration_seconds,
            "text_length": item.text_length,
            "error_message": item.error_message,
        }))

    activities.sort(key=lambda item: item.get("created_at") or "", reverse=True)
    return activities[:limit]


def categorize_activities(activities: list[dict]) -> dict:
    categories = {
        "admin": [],
        "training": [],
        "assignment": [],
        "video": [],
        "face": [],
        "speech": [],
        "audit": [],
        "profile": [],
    }
    for item in activities:
        categories.setdefault(item.get("category") or "profile", []).append(item)
    return categories


def serialize_usage_account(db: Session, user: models.User, *, include_activities: bool = False) -> dict:
    stats = account_usage_stats(db, user)
    changed_at = account_usage_changed_at(db, user)
    total_actions = sum(int(value or 0) for value in stats.values())
    data = {
        "id": user.id,
        "username": user.username,
        "role": user.role,
        "display_name": user.display_name,
        "real_name": user.real_name,
        "unit": user.unit,
        "department": user.department,
        "created_at": user.created_at.isoformat() if user.created_at else None,
        "last_login_at": user.last_login_at.isoformat() if user.last_login_at else None,
        "profile": account_profile(db, user),
        "stats": stats,
        "total_action_count": total_actions,
        "usage_version": timestamp_token(changed_at),
        "usage_changed_at": changed_at.isoformat() if changed_at else None,
    }
    if include_activities:
        activities = account_recent_activities(db, user, limit=120)
        data["activities"] = activities
        data["activity_categories"] = categorize_activities(activities)
    else:
        activities = account_recent_activities(db, user, limit=8)
        data["recent_activities"] = activities[:3]
        data["activity_categories"] = categorize_activities(activities)
    return data


def delete_student_owned_data(db: Session, user_id: int) -> None:
    delete_media_assets(db, owner_type="user", owner_key=user_id, asset_kind="avatar")
    delete_media_assets(db, owner_type="face_profile", owner_key=user_id)
    training_session_ids = [
        row[0]
        for row in db.query(models.TrainingSession.id)
        .filter(models.TrainingSession.user_id == user_id)
        .all()
    ]
    if training_session_ids:
        artifact_ids = [
            row[0]
            for row in db.query(models.TrainingSessionArtifact.id)
            .filter(models.TrainingSessionArtifact.session_id.in_(training_session_ids))
            .all()
        ]
        for artifact_id in artifact_ids:
            delete_media_assets(db, owner_type="training_session_artifact", owner_key=artifact_id)
        delete_optional_table_rows(db, "face_verification_sessions", "session_id", training_session_ids)
        delete_optional_table_rows(db, "multimodal_events", "session_id", training_session_ids)
        delete_optional_table_rows(db, "multimodal_session_metrics", "session_id", training_session_ids)
        db.query(models.FaceVerificationEvent).filter(
            models.FaceVerificationEvent.session_id.in_(training_session_ids)
        ).delete(synchronize_session=False)
        db.query(models.Message).filter(models.Message.session_id.in_(training_session_ids)).delete(synchronize_session=False)
        db.query(models.TrainingSessionArtifact).filter(models.TrainingSessionArtifact.session_id.in_(training_session_ids)).delete(synchronize_session=False)
        db.query(models.AssignmentSubmission).filter(models.AssignmentSubmission.training_session_id.in_(training_session_ids)).update(
            {models.AssignmentSubmission.training_session_id: None},
            synchronize_session=False,
        )
        db.query(models.TrainingSession).filter(models.TrainingSession.id.in_(training_session_ids)).delete(synchronize_session=False)

    video_session_ids = [
        row[0]
        for row in db.query(models.VideoTrainingSession.id)
        .filter(models.VideoTrainingSession.user_id == user_id)
        .all()
    ]
    delete_video_training_sessions(db, video_session_ids)

    db.query(models.AssignmentSubmission).filter(models.AssignmentSubmission.user_id == user_id).delete(synchronize_session=False)
    db.query(models.AssignmentStudentOverride).filter(models.AssignmentStudentOverride.user_id == user_id).delete(synchronize_session=False)
    db.query(models.ClassMembership).filter(models.ClassMembership.user_id == user_id).delete(synchronize_session=False)
    delete_optional_table_rows(db, "face_verification_sessions", "student_id", [user_id])
    delete_optional_table_rows(db, "multimodal_events", "student_id", [user_id])
    delete_optional_table_rows(db, "multimodal_session_metrics", "student_id", [user_id])
    db.query(models.FaceVerificationEvent).filter(models.FaceVerificationEvent.student_id == user_id).delete(synchronize_session=False)
    db.query(models.FaceProfile).filter(models.FaceProfile.student_id == user_id).delete(synchronize_session=False)
    db.query(models.SpeechUsageLog).filter(models.SpeechUsageLog.user_id == user_id).delete(synchronize_session=False)


def delete_user_references(db: Session, user_id: int) -> None:
    db.query(models.OpsAuditLog).filter(
        (models.OpsAuditLog.actor_id == user_id) | (models.OpsAuditLog.target_user_id == user_id)
    ).delete(synchronize_session=False)
    db.query(models.SpeechUsageLog).filter(models.SpeechUsageLog.user_id == user_id).delete(synchronize_session=False)


@router.get("/accounts", response_model=list[schemas.OpsAccountOverview])
def list_accounts(
    role: str | None = None,
    group: str | None = None,
    keyword: str | None = None,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    query = db.query(models.User)
    normalized_role = str(role or "").strip()
    if normalized_role:
        if normalized_role not in VISIBLE_ROLES:
            raise HTTPException(status_code=400, detail="不支持的账号角色")
        query = query.filter(models.User.role == normalized_role)
    else:
        query = query.filter(models.User.role.in_(VISIBLE_ROLES))

    search_text = str(keyword or "").strip()
    if search_text:
        query = query.filter(models.User.username.contains(search_text))

    group_text = str(group or "").strip()
    if group_text:
        query = query.filter(models.User.account_group == group_text)

    users = query.order_by(models.User.account_group.asc(), models.User.role.asc(), models.User.username.asc()).all()
    return [serialize_account(user) for user in users]


@router.get("/account-groups")
def list_account_groups(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    registry_rows = db.query(models.AccountGroup).order_by(models.AccountGroup.name.asc()).all()
    user_rows = (
        db.query(models.User.account_group, func.count(models.User.id))
        .filter(models.User.role.in_(VISIBLE_ROLES))
        .group_by(models.User.account_group)
        .all()
    )
    counts = {str(name or "").strip(): int(count or 0) for name, count in user_rows if str(name or "").strip()}
    ungrouped_count = sum(int(count or 0) for name, count in user_rows if not str(name or "").strip())
    groups = []
    seen: set[str] = set()
    for row in registry_rows:
        group_name = normalize_account_group_name(row.name)
        if not group_name or group_name in seen:
            continue
        seen.add(group_name)
        groups.append({
            "id": row.id,
            "name": group_name,
            "description": row.description,
            "count": counts.get(group_name, 0),
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "updated_at": row.updated_at.isoformat() if row.updated_at else None,
        })
    for group_name, count in sorted(counts.items(), key=lambda item: item[0]):
        if group_name not in seen:
            groups.append({
                "id": None,
                "name": group_name,
                "description": None,
                "count": count,
                "created_at": None,
                "updated_at": None,
            })
            seen.add(group_name)
    return {"groups": groups, "ungrouped_count": ungrouped_count}


@router.post("/account-groups", response_model=schemas.AccountGroupOverview)
def create_account_group(
    payload: schemas.AccountGroupCreate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    name = normalize_account_group_name(payload.name)
    if not name:
        raise HTTPException(status_code=400, detail="请输入分组名称")
    existing = db.query(models.AccountGroup).filter(models.AccountGroup.name == name).first()
    if existing:
        raise HTTPException(status_code=400, detail="分组已存在")
    group = ensure_account_group_registry(db, name, description=payload.description)
    db.flush()
    write_ops_audit(db, actor=current_user, action="create_account_group", detail={"name": name, "description": group.description})
    db.commit()
    db.refresh(group)
    return schemas.AccountGroupOverview(
        id=group.id,
        name=group.name,
        description=group.description,
        count=0,
        created_at=group.created_at,
        updated_at=group.updated_at,
    )


@router.get("/accounts/usage")
def list_account_usage(
    role: str | None = None,
    group: str | None = None,
    keyword: str | None = None,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    query = db.query(models.User).filter(models.User.role.in_(VISIBLE_ROLES))
    normalized_role = str(role or "").strip()
    if normalized_role:
        if normalized_role not in VISIBLE_ROLES:
            raise HTTPException(status_code=400, detail="不支持的账号角色")
        query = query.filter(models.User.role == normalized_role)

    search_text = str(keyword or "").strip()
    if search_text:
        query = query.filter(models.User.username.contains(search_text))
    group_text = str(group or "").strip()
    if group_text:
        query = query.filter(models.User.account_group == group_text)

    users = query.order_by(models.User.account_group.asc(), models.User.role.asc(), models.User.username.asc()).all()
    return [serialize_usage_account(db, user) for user in users]


@router.get("/accounts/usage-versions")
def list_account_usage_versions(
    role: str | None = None,
    group: str | None = None,
    keyword: str | None = None,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    query = db.query(models.User).filter(models.User.role.in_(VISIBLE_ROLES))
    normalized_role = str(role or "").strip()
    if normalized_role:
        if normalized_role not in VISIBLE_ROLES:
            raise HTTPException(status_code=400, detail="不支持的账号角色")
        query = query.filter(models.User.role == normalized_role)

    search_text = str(keyword or "").strip()
    if search_text:
        query = query.filter(models.User.username.contains(search_text))
    group_text = str(group or "").strip()
    if group_text:
        query = query.filter(models.User.account_group == group_text)

    users = query.order_by(models.User.account_group.asc(), models.User.role.asc(), models.User.username.asc()).all()
    accounts = [serialize_usage_version(db, user) for user in users]
    global_changed_at = max(
        (datetime.fromisoformat(item["changed_at"]) for item in accounts if item.get("changed_at")),
        default=None,
    )
    return {
        "global_version": timestamp_token(global_changed_at),
        "changed_at": global_changed_at.isoformat() if global_changed_at else None,
        "accounts": accounts,
    }


@router.get("/accounts/{account_id}/usage")
def get_account_usage(
    account_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    user = ensure_account(db, account_id)
    return serialize_usage_account(db, user, include_activities=True)


@router.post("/accounts", response_model=schemas.OpsAccountOverview)
def create_account(
    payload: schemas.OpsAccountCreate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    username = str(payload.username or "").strip()
    password = str(payload.password or "").strip()
    role = str(payload.role or "").strip()
    if not username:
        raise HTTPException(status_code=400, detail="请输入账号")
    if len(password) < 6:
        raise HTTPException(status_code=400, detail="密码至少 6 位")
    if role not in MANAGED_ROLES:
        raise HTTPException(status_code=400, detail="维护端只能创建管理端账号或学员端账号")
    if db.query(models.User).filter(models.User.username == username).first():
        raise HTTPException(status_code=400, detail="账号已存在")

    account_group = ensure_account_group_registry(db, getattr(payload, "account_group", None))
    user = models.User(username=username, hashed_password=hash_password(password), role=role, account_group=account_group.name if account_group else None)
    apply_profile_fields(user, payload)
    db.add(user)
    db.flush()
    write_ops_audit(db, actor=current_user, action="create_account", target_user=user, detail={
        "username": username,
        "role": role,
        "account_group": user.account_group,
    })
    db.commit()
    db.refresh(user)
    return serialize_account(user)


@router.post("/accounts/batch", response_model=schemas.BatchStudentCreateResponse)
def batch_create_student_accounts(
    payload: schemas.BatchStudentCreateRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    password = str(payload.password or "").strip()
    account_group = clean_text(getattr(payload, "account_group", None), 80)
    if not password:
        raise HTTPException(status_code=400, detail="请输入初始密码")
    usernames = build_username_range(payload.template, int(payload.start_no), int(payload.end_no))
    created_usernames: list[str] = []
    skipped_usernames: list[str] = []
    for username in usernames:
        if db.query(models.User).filter(models.User.username == username).first():
            skipped_usernames.append(username)
            continue
        ensure_account_group_registry(db, account_group)
        db.add(models.User(username=username, hashed_password=hash_password(password), role="student", account_group=account_group))
        created_usernames.append(username)
    if created_usernames:
        write_ops_audit(db, actor=current_user, action="batch_create_students", detail={
            "template": payload.template,
            "start_no": payload.start_no,
            "end_no": payload.end_no,
            "created_count": len(created_usernames),
            "skipped_count": len(skipped_usernames),
            "account_group": account_group,
        })
        db.commit()
    return schemas.BatchStudentCreateResponse(
        created_count=len(created_usernames),
        skipped_count=len(skipped_usernames),
        created_usernames=created_usernames,
        skipped_usernames=skipped_usernames,
    )


@router.post("/accounts/import/preview", response_model=schemas.OpsAccountImportPreviewResponse)
async def preview_import_accounts(
    file: UploadFile = File(...),
    target_role: str | None = Form(default="student"),
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="导入文件为空")
    records = parse_import_file(file.filename or "", content)
    if not records:
        raise HTTPException(status_code=400, detail="未读取到可导入账号")
    records = apply_target_role(records, target_role)
    return build_import_preview(file.filename or "accounts", records, db)


@router.post("/accounts/import/commit", response_model=schemas.OpsAccountImportCommitResponse)
def commit_import_accounts(
    payload: schemas.OpsAccountImportCommitRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    preview_items = validate_import_records([item.model_dump() for item in payload.accounts], db)
    created_usernames: list[str] = []
    skipped_usernames: list[str] = []
    failed_items: list[schemas.OpsAccountImportPreviewItem] = []

    for item in preview_items:
        if item.status != "ready":
            if "账号已存在" in item.errors:
                skipped_usernames.append(item.username)
            else:
                failed_items.append(item)
            continue
        user = models.User(
            username=item.username,
            hashed_password=hash_password(item.password or "123456"),
            role=item.role,
            account_group=item.account_group,
        )
        for field, max_length in PROFILE_FIELD_LIMITS.items():
            setattr(user, field, clean_text(getattr(item, field), max_length))
        db.add(user)
        db.flush()
        ensure_account_group_registry(db, item.account_group)
        write_ops_audit(db, actor=current_user, action="import_create_account", target_user=user, detail={
            "username": item.username,
            "role": item.role,
            "account_group": item.account_group,
            "row_number": item.row_number,
        })
        created_usernames.append(item.username)

    if created_usernames:
        db.commit()

    return schemas.OpsAccountImportCommitResponse(
        created_count=len(created_usernames),
        skipped_count=len(skipped_usernames),
        failed_count=len(failed_items),
        created_usernames=created_usernames,
        skipped_usernames=skipped_usernames,
        failed_items=failed_items,
    )


@router.post("/accounts/batch-delete/", response_model=schemas.OpsAccountBatchDeleteResponse, include_in_schema=False)
@router.post("/accounts/batch_delete", response_model=schemas.OpsAccountBatchDeleteResponse, include_in_schema=False)
@router.post("/accounts/batch-delete", response_model=schemas.OpsAccountBatchDeleteResponse)
def batch_delete_accounts(
    payload: schemas.OpsAccountBatchDeleteRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    account_ids = [int(account_id) for account_id in payload.account_ids if int(account_id or 0) > 0]
    if not account_ids:
        raise HTTPException(status_code=400, detail="请至少选择一个账号")

    deleted_usernames: list[str] = []
    skipped_usernames: list[str] = []
    for account_id in account_ids:
        user = db.query(models.User).filter(models.User.id == account_id).first()
        if not user:
            continue
        if user.id == current_user.id:
            skipped_usernames.append(user.username)
            continue
        if user.role == "maintainer":
            skipped_usernames.append(user.username)
            continue
        deleted_usernames.append(user.username)
        write_ops_audit(db, actor=current_user, action="delete_account", detail={
            "target_user_id": user.id,
            "username": user.username,
            "role": user.role,
            "batch": True,
        })
        if user.role in MANAGED_ROLES:
            delete_student_owned_data(db, user.id)
        if user.role == "admin":
            delete_admin_owned_data(db, user.id)
        delete_user_references(db, user.id)
        db.delete(user)

    if deleted_usernames:
        db.commit()

    return schemas.OpsAccountBatchDeleteResponse(
        deleted_count=len(deleted_usernames),
        skipped_count=len(skipped_usernames),
        deleted_usernames=deleted_usernames,
        skipped_usernames=skipped_usernames,
    )


@router.patch("/accounts/{account_id}", response_model=schemas.OpsAccountOverview)
def update_account(
    account_id: int,
    payload: schemas.OpsAccountUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    user = ensure_account(db, account_id)
    if user.role == "maintainer" and user.id != current_user.id:
        raise HTTPException(status_code=400, detail="不能在此控制台编辑其他维护账号")

    if "username" in payload.model_fields_set and payload.username is not None:
        username = str(payload.username or "").strip()
        if not username:
            raise HTTPException(status_code=400, detail="请输入账号")
        duplicate = db.query(models.User).filter(models.User.username == username, models.User.id != user.id).first()
        if duplicate:
            raise HTTPException(status_code=400, detail="账号已存在")
        user.username = username

    if "role" in payload.model_fields_set and payload.role is not None:
        role = str(payload.role or "").strip()
        if user.id == current_user.id:
            raise HTTPException(status_code=400, detail="不能修改当前维护账号的角色")
        if role not in MANAGED_ROLES:
            raise HTTPException(status_code=400, detail="角色只能是管理端账号或学员端账号")
        user.role = role

    apply_profile_fields(user, payload)
    if "account_group" in payload.model_fields_set:
        account_group = ensure_account_group_registry(db, getattr(payload, "account_group", None))
        user.account_group = account_group.name if account_group else None
    user.updated_at = datetime.utcnow()
    db.add(user)
    write_ops_audit(db, actor=current_user, action="update_account", target_user=user, detail={
        "username": user.username,
        "role": user.role,
        "fields": sorted(payload.model_fields_set),
    })
    db.commit()
    db.refresh(user)
    return serialize_account(user)


@router.post("/account-groups/delete/", response_model=schemas.AccountGroupDeleteResponse, include_in_schema=False)
@router.post("/account-groups/delete", response_model=schemas.AccountGroupDeleteResponse)
def delete_account_group(
    payload: schemas.AccountGroupDeleteRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    group_name = normalize_account_group_name(payload.name)
    if not group_name:
        raise HTTPException(status_code=400, detail="请输入分组名称")
    group = db.query(models.AccountGroup).filter(models.AccountGroup.name == group_name).first()
    affected_users = db.query(models.User).filter(models.User.account_group == group_name).all()
    if not group and not affected_users:
        raise HTTPException(status_code=404, detail="分组不存在")
    deleted_accounts_count = 0
    skipped_usernames: list[str] = []
    if payload.delete_accounts:
        for user in affected_users:
            if user.role == "maintainer" or user.id == current_user.id:
                skipped_usernames.append(user.username)
                continue
            if user.role in MANAGED_ROLES:
                delete_student_owned_data(db, user.id)
            if user.role == "admin":
                delete_admin_owned_data(db, user.id)
            delete_user_references(db, user.id)
            write_ops_audit(db, actor=current_user, action="delete_account", detail={
                "target_user_id": user.id,
                "username": user.username,
                "role": user.role,
                "group_deleted": group_name,
            })
            db.delete(user)
            deleted_accounts_count += 1
    else:
        for user in affected_users:
            user.account_group = None
            db.add(user)

    if group:
        db.delete(group)
    write_ops_audit(db, actor=current_user, action="delete_account_group", detail={
        "name": group_name,
        "delete_accounts": bool(payload.delete_accounts),
        "deleted_accounts_count": deleted_accounts_count,
        "skipped_accounts_count": len(skipped_usernames),
    })
    db.commit()
    return schemas.AccountGroupDeleteResponse(
        deleted_group=group_name,
        deleted_accounts_count=deleted_accounts_count,
        skipped_accounts_count=len(skipped_usernames),
        skipped_usernames=skipped_usernames,
    )


@router.post("/accounts/{account_id}/reset-password")
def reset_account_password(
    account_id: int,
    payload: schemas.OpsPasswordResetRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    user = ensure_account(db, account_id)
    if user.role == "maintainer" and user.id != current_user.id:
        raise HTTPException(status_code=400, detail="不能在此控制台重置其他维护账号密码")
    new_password = str(payload.new_password or "").strip()
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="密码至少 6 位")
    user.hashed_password = hash_password(new_password)
    user.updated_at = datetime.utcnow()
    db.add(user)
    write_ops_audit(db, actor=current_user, action="reset_password", target_user=user, detail={"username": user.username})
    db.commit()
    return {"success": True}


@router.delete("/accounts/{account_id}")
def delete_account(
    account_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    user = ensure_account(db, account_id)
    if user.id == current_user.id:
        raise HTTPException(status_code=400, detail="不能删除当前登录的维护账号")
    if user.role == "maintainer":
        raise HTTPException(status_code=400, detail="不能在此控制台删除维护账号")
    if user.role not in MANAGED_ROLES:
        raise HTTPException(status_code=400, detail="不支持的账号角色")

    deleted_username = user.username
    deleted_role = user.role
    write_ops_audit(db, actor=current_user, action="delete_account", detail={
        "target_user_id": user.id,
        "username": deleted_username,
        "role": deleted_role,
    })
    if user.role in MANAGED_ROLES:
        delete_student_owned_data(db, user.id)
    if user.role == "admin":
        delete_admin_owned_data(db, user.id)
    delete_user_references(db, user.id)

    db.delete(user)
    db.commit()
    return {"success": True, "deleted_id": account_id, "deleted_username": deleted_username}


def _parse_ops_json(value: str | None) -> dict:
    try:
        parsed = json.loads(value or "{}")
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


def _serialize_system_issue(item: models.OpsIssueRecord) -> dict:
    return {
        "id": item.id,
        "category": item.category,
        "severity": item.severity,
        "status": item.status,
        "source": item.source,
        "case_id": item.case_id,
        "workflow_run_id": item.workflow_run_id,
        "title": item.title,
        "detail": item.detail,
        "metadata": _parse_ops_json(item.metadata_json),
        "created_at": item.created_at.isoformat() if item.created_at else None,
        "updated_at": item.updated_at.isoformat() if item.updated_at else None,
    }


@router.get("/system-issues")
def list_system_issues(
    category: str | None = None,
    status: str | None = None,
    provider: str | None = None,
    limit: int = 100,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    query = db.query(models.OpsIssueRecord)
    if category:
        query = query.filter(models.OpsIssueRecord.category == category)
    if status:
        query = query.filter(models.OpsIssueRecord.status == status)
    if provider:
        query = query.filter(models.OpsIssueRecord.metadata_json.contains(f'"provider": "{provider}"'))
    rows = query.order_by(models.OpsIssueRecord.created_at.desc()).limit(max(1, min(limit, 300))).all()
    return [_serialize_system_issue(item) for item in rows]


@router.get("/ai-workflow-runs")
def list_ai_workflow_runs(
    stage: str | None = None,
    provider: str | None = None,
    limit: int = 100,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    query = db.query(models.AIWorkflowRun)
    if stage:
        query = query.filter(models.AIWorkflowRun.stage == stage)
    if provider:
        query = query.filter((models.AIWorkflowRun.primary_provider == provider) | (models.AIWorkflowRun.final_provider == provider))
    rows = query.order_by(models.AIWorkflowRun.created_at.desc()).limit(max(1, min(limit, 300))).all()
    return [{
        "id": row.id, "correlation_id": row.correlation_id, "case_id": row.case_id, "stage": row.stage,
        "status": row.status, "primary_provider": row.primary_provider, "final_provider": row.final_provider,
        "model": row.model, "attempt_count": row.attempt_count, "switched_provider": row.switched_provider,
        "used_rule_fallback": row.used_rule_fallback, "error_code": row.error_code, "error_summary": row.error_summary,
        "trace": _parse_ops_json(row.trace_json), "created_at": row.created_at.isoformat() if row.created_at else None,
    } for row in rows]


@router.patch("/system-issues/{issue_id}")
def update_system_issue(
    issue_id: int,
    payload: dict = Body(...),
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(require_maintainer_user),
):
    issue = db.query(models.OpsIssueRecord).filter(models.OpsIssueRecord.id == issue_id).first()
    if not issue:
        raise HTTPException(status_code=404, detail="异常记录不存在")
    next_status = str(payload.get("status") or "").strip()
    if next_status not in {"pending", "acknowledged", "resolved", "ignored"}:
        raise HTTPException(status_code=400, detail="不支持的异常状态")
    issue.status = next_status
    db.add(issue)
    write_ops_audit(db, actor=current_user, action="update_system_issue", detail={"issue_id": issue.id, "status": next_status})
    db.commit()
    db.refresh(issue)
    return _serialize_system_issue(issue)
