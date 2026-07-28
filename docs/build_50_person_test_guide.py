from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "AI虚拟警情处置模拟训练平台_50人双端全员测试使用说明_v1.1.docx"

BLUE = "165DFF"
NAVY = "16324F"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FB"
LIGHT_GRAY = "F2F4F7"
GREEN = "166534"
GREEN_BG = "ECFDF3"
GOLD = "7A5A00"
GOLD_BG = "FFF8E1"
RED = "9B1C1C"
RED_BG = "FEF2F2"
WHITE = "FFFFFF"
DXA_WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D7DEE8", size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, cn="Microsoft YaHei", latin="Calibri"):
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), cn)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=INK, bold=False):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, MUTED)


def add_para(doc, text="", *, size=11, color=INK, bold=False, italic=False, align=None,
             before=0, after=6, line=1.25, keep=False, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, *, size=11, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    for part in parts:
        text = part[0]
        kwargs = part[1] if len(part) > 1 else {}
        r = p.add_run(text)
        set_run_font(r, size=kwargs.get("size", size), color=kwargs.get("color", INK),
                     bold=kwargs.get("bold", False), italic=kwargs.get("italic", False))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="Bullet Compact" if level == 0 else "Bullet Compact 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="Number Compact")
    p.add_run(text)
    return p


def add_check(doc, text, checked=False):
    marker = "☒" if checked else "☐"
    return add_rich_para(doc, [
        (marker + "  ", {"bold": True, "color": BLUE}),
        (text, {}),
    ], after=4)


def add_callout(doc, label, text, kind="info"):
    palette = {
        "info": (BLUE, LIGHTER_BLUE),
        "success": (GREEN, GREEN_BG),
        "warning": (GOLD, GOLD_BG),
        "danger": (RED, RED_BG),
    }
    accent, fill = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [DXA_WIDTH], 120)
    set_table_borders(table, color=fill, size=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}  ")
    set_run_font(r, 10.5, accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    add_para(doc, "", after=2)


def add_table(doc, headers, rows, widths, aligns=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths, 120)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(str(value))
        set_run_font(r, font_size, NAVY, bold=True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if aligns and i < len(aligns):
                p.alignment = aligns[i]
            r = p.add_run(str(value))
            set_run_font(r, font_size, INK)
    add_para(doc, "", after=3)
    return table


def setup_numbering(doc):
    numbering = doc.part.numbering_part.element
    max_abstract = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] or [0])
    max_num = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0])

    def make_num(fmt, text, left, hanging, font="Microsoft YaHei"):
        nonlocal max_abstract, max_num
        max_abstract += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(max_abstract))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:eastAsia"), font)
        r_pr.append(r_fonts)
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)
        max_num += 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(max_num))
        abstract_id = OxmlElement("w:abstractNumId")
        abstract_id.set(qn("w:val"), str(max_abstract))
        num.append(abstract_id)
        numbering.append(num)
        return max_num

    bullet_id = make_num("bullet", "•", 540, 270)
    bullet2_id = make_num("bullet", "–", 900, 270)
    number_id = make_num("decimal", "%1.", 540, 270)

    for name, num_id in (("Bullet Compact", bullet_id), ("Bullet Compact 2", bullet2_id), ("Number Compact", number_id)):
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, 11, INK)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        p_pr = style._element.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        p_pr.append(num_pr)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, NAVY, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.1

    setup_numbering(doc)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AI虚拟警情处置模拟训练平台  |  50人测试使用说明")
    set_run_font(r, 8.5, MUTED, bold=True)
    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_cover(doc):
    add_para(doc, "测试执行手册", size=11, color=BLUE, bold=True, before=18, after=8)
    add_para(doc, "AI虚拟警情处置模拟训练平台", size=27, color=NAVY, bold=True, after=8, line=1.05)
    add_para(doc, "50 人双端全员测试使用说明", size=18, color=BLUE, bold=True, after=14)
    add_para(doc, "面向管理员、学员与测试组织者的完整操作、验收与反馈指南", size=12.5, color=MUTED, after=24)

    add_callout(
        doc,
        "本次测试目标",
        "用 50 名真实用户验证“管理员建案发布—学员训练—评估复盘—管理端查看结果”的双端完整闭环，同时观察并发稳定性、AI 对话质量、评估有效性、数据隔离与首次使用易用性。",
        "info",
    )

    add_table(doc, ["文档项", "内容"], [
        ["适用人数", "50 名测试人员；建议另设 1 名总协调、1 名技术保障、1—2 名教官观察员"],
        ["建议时长", "单人 150—180 分钟；其中统一并发窗口 10—15 分钟"],
        ["必测范围", "管理端内容配置与发布、知识/班级等管理；学员端训练、评估、历史；跨端状态、并发和权限"],
        ["账号要求", "每名测试人员使用一组独立管理员账号和学员账号，按个人测试编号创建数据"],
        ["测试环境", "平台地址、账号、班级邀请码和指定案件由组织者在测试前统一发放"],
        ["文档版本", "v1.0  |  编制日期：2026 年 7 月 14 日"],
    ], [1800, 7560], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 10)

    add_heading(doc, "开始前先看", 2)
    add_check(doc, "使用分配给本人的独立账号，不要多人共用同一账号。")
    add_check(doc, "训练内容使用模拟案情；不要输入真实姓名、身份证号、电话、住址或未公开案件信息。")
    add_check(doc, "遇到问题先截图并记录时间、页面和操作步骤，不要反复提交同一请求。")
    add_check(doc, "除被分配“异常恢复”任务外，不要清除浏览器缓存、删除训练记录或强制关闭页面。")
    add_para(doc, "本说明中的验收阈值为本轮 50 人试测的建议值。若项目负责人另行发布更严格标准，以最新通知为准。", size=9.5, color=MUTED, italic=True, before=8, after=0)
    doc.add_page_break()


def add_quick_card(doc):
    add_heading(doc, "一、测试人员一页速览", 1)
    add_callout(doc, "你需要完成", "先使用管理员账号完成建案、角色/场景/考察点配置并发布；再使用配对学员账号完成至少 1 次不少于 8 轮的训练、评估和历史复盘；最后回到管理端核对结果。", "success")

    add_table(doc, ["阶段", "你要做什么", "完成标志"], [
        ["1. 管理员登录", "用管理员账号进入管理端并确认权限", "看到仪表盘；无反复跳转或权限错误"],
        ["2. 建案发布", "创建案件、角色、场景和考察点，保存并发布", "本人编号的已发布场景可用于学员训练"],
        ["3. 管理验证", "完成知识库、班级/作业、学员与其他已启用模块的基本操作", "保存、检索、状态显示正确"],
        ["4. 学员训练", "切换配对学员账号，找到本人发布内容并完成 ≥8 轮对话", "对话可连续进行，AI 角色能回应当前问题"],
        ["5. 生成报告", "结束训练/提交作业并生成评估", "报告成功显示且内容与本次对话相关"],
        ["6. 跨端核对", "学员历史复盘后回到管理端，查看状态或作业完成情况", "发布、训练、报告与管理结果一致"],
        ["7. 提交反馈", "填写双端评分与问题记录；严重问题立即报告", "问题可复现、信息完整、带截图/时间"],
    ], [1050, 5140, 3170], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 9.2)

    add_heading(doc, "建议测试口径", 2)
    add_bullet(doc, "“一轮有效交流”指学员发送一条有明确处置目的的内容，并收到 AI 针对性回应；问候、重复发送、无意义字符不计。")
    add_bullet(doc, "测试 AI 不是背固定台词。请按真实执法思路自然表达，但不得把真实敏感案情输入平台。")
    add_bullet(doc, "如 AI 正在生成回复或报告，请等待页面结束加载。超过本说明规定的观察时间，再按故障流程记录。")
    add_bullet(doc, "如果某功能未由组织者开放（例如视频实训），在反馈表标记“未配置/不适用”；其余管理端与学员端主链路均为全员必测。")


def add_scope_and_groups(doc):
    add_heading(doc, "二、要测试什么", 1)
    add_para(doc, "本轮测试分为“全员必测”和“小组专项”两层。全员数据用于判断主流程是否可用；专项测试用于扩大浏览器、输入方式、恢复路径和多模态能力的覆盖。")

    add_heading(doc, "2.1 全员必测范围", 2)
    for text in [
        "账号登录、角色识别与退出重登。",
        "训练大厅浏览、状态显示、案件难度筛选和指定场景进入。",
        "如启用班级作业：邀请码加入、公告/截止时间查看、作业开始与完成状态更新。",
        "文本方式完成不少于 8 轮有效警情处置对话。",
        "接警简报、训练阶段/考察进度、角色发言与页面反馈是否清晰。",
        "结束训练并生成评估报告；检查总评、场景考察点、扣分依据和改进建议。",
        "训练历史、对话记录、评估报告的保存与再次打开。",
        "页面响应、错误提示、信息安全和整体易用性。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.2 50 人分组建议", 2)
    add_table(doc, ["小组", "人数", "所有人共同任务", "额外专项关注"], [
        ["A 组", "10", "完整主流程", "首次使用易用性：不接受现场口头教学，记录所有不理解的文字、按钮和步骤"],
        ["B 组", "10", "完整主流程", "语音与多角色：至少 2 轮语音输入，观察识别、打断、角色切换与上下文衔接"],
        ["C 组", "10", "完整主流程", "中断恢复：完成 3 轮后正常离开/刷新，再从训练历史继续；核对消息与进度"],
        ["D 组", "10", "完整主流程", "评估一致性：同一案件复训一次，比较考察点、扣分依据与改进建议是否可解释"],
        ["E 组", "10", "完整主流程", "班级作业与视频实训（若已配置）；验证摄像头/麦克风、节点、回放与报告"],
    ], [900, 700, 2200, 5560], [WD_ALIGN_PARAGRAPH.CENTER] * 2 + [WD_ALIGN_PARAGRAPH.LEFT] * 2, 8.8)

    add_callout(doc, "统一并发窗口", "建议在测试开始后的第 15—30 分钟组织一次集中操作：50 人同时在线，至少 40 人在 5 分钟内提交对话，至少 30 人在同一分钟内发起 AI 回复。技术保障人员同步观察接口错误、响应时间和服务资源。", "warning")


def add_preparation(doc):
    add_heading(doc, "三、测试前准备", 1)
    add_heading(doc, "3.1 组织者准备清单（测试前 1 天完成）", 2)
    for text in [
        "确认测试环境地址可从 50 台设备访问，HTTPS、域名、网络策略和浏览器证书均正常。",
        "创建 50 组独立账号（管理员 + 学员），形成账号—人员—编号对应表；不要使用生产默认密码。",
        "为每个编号准备脱敏的案件材料、角色信息、场景要求、至少 3 条考察点与知识库小文档；先用一组账号完整走通。",
        "提前确认 50 名管理员可以独立创建测试数据，向所有人强调只操作本人【Txx】编号的数据；如采用班级作业，准备相应邀请码。",
        "如测试视频实训，确认视频已发布、节点配置完整、摄像头/麦克风权限策略可用，并准备无摄像头设备的降级安排。",
        "备份测试前数据；确认本轮数据可与正式数据区分，约定测试结束后的保留或清理策略。",
        "建立唯一问题收集渠道，指定总协调人、教官观察员和技术保障人，公布紧急联系方式。",
        "准备统一计时点：测试开始、并发窗口开始、并发窗口结束、测试结束。",
    ]:
        add_check(doc, text)

    add_heading(doc, "3.2 测试人员设备与环境", 2)
    add_table(doc, ["项目", "要求"], [
        ["设备", "优先使用电脑；屏幕建议 1366×768 及以上。视频专项需摄像头和麦克风。"],
        ["浏览器", "优先最新版 Chrome 或 Edge；不要使用无痕模式，除非组织者要求验证无痕兼容性。"],
        ["网络", "使用稳定网络；测试中途不要主动切换 Wi-Fi/热点，异常恢复组除外。"],
        ["声音环境", "语音/视频专项建议佩戴耳机，在相对安静空间测试，避免 50 台设备相互串音。"],
        ["资料", "管理员账号、学员账号、个人测试编号、测试地址、脱敏案情包、班级邀请码（如有）、反馈入口。"],
        ["隐私", "不得输入真实个人信息、未公开案件材料或其他敏感数据；截图前检查是否含账号密码。"],
    ], [1500, 7860], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.5)

    add_heading(doc, "3.3 现场信息栏（由组织者填写）", 2)
    add_table(doc, ["项目", "现场填写"], [
        ["测试日期/时段", "________________________________________"],
        ["平台访问地址", "________________________________________"],
        ["双账号命名范围", "________________________________________"],
        ["指定案件/场景", "________________________________________"],
        ["班级邀请码（如有）", "________________________________________"],
        ["问题反馈入口", "________________________________________"],
        ["现场联系人", "________________________________________"],
    ], [2200, 7160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.5)


def add_steps(doc):
    add_heading(doc, "五、怎么测试：学员端（全员必测）", 1)
    add_callout(doc, "建议用时", "登录与找任务 10 分钟；主训练 25—35 分钟；报告与历史 15 分钟；问卷与缺陷整理 10—20 分钟。", "info")

    add_heading(doc, "步骤 1：登录并确认身份", 2)
    add_number(doc, "打开组织者提供的平台地址，进入“统一身份登录”。")
    add_number(doc, "输入本人账号和密码，点击“登录”；不要保存或转发密码截图。")
    add_number(doc, "确认登录后进入学员端，而不是管理员端或维护端。")
    add_number(doc, "记录登录是否一次成功、页面是否在 3 秒左右可操作、错误提示是否能看懂。")
    add_callout(doc, "预期结果", "显示登录成功并进入学员首页；账号角色正确；刷新页面后仍保持有效登录状态；错误账号应得到明确提示而非空白页。", "success")

    add_heading(doc, "步骤 2：找到指定训练", 2)
    add_number(doc, "在训练大厅查看可用案件，并尝试一次类型或难度筛选。")
    add_number(doc, "确认案件标题、场景名称、难度、简介和训练状态可识别。")
    add_number(doc, "若组织者采用班级作业：进入“我的班级/班级作业”，输入邀请码，查看公告、截止时间、完成标准，再开始指定场景。")
    add_number(doc, "点击“开始训练”；如已有未完成记录，应显示“继续训练”并恢复原会话。")
    add_callout(doc, "预期结果", "已发布的指定任务可见；筛选结果正确；班级邀请码能加入一次且不会重复入班；开始训练后能进入正确案件和场景。", "success")

    add_heading(doc, "步骤 3：阅读接警简报并检查训练页面", 2)
    add_number(doc, "进入训练后阅读接警简报/案件信息，核对背景、地点、人物与组织者指定内容是否一致。")
    add_number(doc, "观察训练阶段或考察进度、对话区、角色信息、输入区和结束训练入口是否清楚。")
    add_number(doc, "如页面提供“今日不再提示”，可关闭弹窗后从页面侧栏再次打开，确认信息仍可查看。")
    add_callout(doc, "预期结果", "简报能正常打开和关闭；核心案情不乱码、不缺失；页面没有遮挡、重叠、按钮不可见或明显错位。", "success")

    add_heading(doc, "步骤 4：完成不少于 8 轮有效对话", 2)
    add_para(doc, "请按真实警情处置逻辑自然交流，不要机械复制下列内容。至少覆盖其中 6 类，并记录 AI 是否逐步、合理地释放信息。")
    add_table(doc, ["对话目标", "可采用的自然表达方向", "重点观察"], [
        ["身份与安全", "表明身份，确认现场是否仍有冲突、伤情、危险物品或紧急风险", "AI 是否理解警察身份；高风险信息是否得到合理回应"],
        ["时间地点", "确认发生时间、具体位置、现场范围和人员是否仍在场", "回答是否具体，前后是否一致"],
        ["人物关系", "分别确认报警人、当事人、证人身份及彼此关系", "多角色是否混淆；称谓和立场是否连贯"],
        ["事件经过", "用开放式问题了解起因、过程、结果，再追问关键矛盾", "是否一次性泄露全部信息；是否能根据追问补充"],
        ["证据线索", "询问监控、照片、录音、聊天记录、物证或现场目击者", "是否出现无根据的事实；线索是否与案情一致"],
        ["情绪与沟通", "对激动、抵触或隐瞒对象进行安抚、告知和边界说明", "情绪/配合变化是否可感知，回应是否像同一角色"],
        ["规范处置", "说明下一步调查、隔离、救助、登记、通知或依法处置安排", "是否出现明显不专业、越权或乱引法条"],
        ["闭环确认", "复述关键信息，确认遗漏，告知后续流程并结束现场交流", "AI 是否响应闭环，评估是否识别到关键动作"],
    ], [1450, 4660, 3250], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.8)

    add_heading(doc, "步骤 5：结束训练并生成评估", 2)
    add_number(doc, "确认已完成至少 8 轮有效交流，点击“结束训练并生成评估”或“提交作业并生成评估”。")
    add_number(doc, "等待系统生成报告；不要连续点击、返回或刷新。超过 60 秒仍无结果时记录问题。")
    add_number(doc, "检查报告编号/时间、总评、场景考察点、扣分发言或依据、改进建议是否与本次对话匹配。")
    add_number(doc, "尝试打印/下载入口（如现场允许），确认不会出现空白报告或错误人员数据。")
    add_callout(doc, "预期结果", "报告生成成功；至少能解释已完成与未完成的考察点；建议具有可操作性；报告只属于当前账号和当前训练记录。", "success")

    add_heading(doc, "步骤 6：训练历史与复盘", 2)
    add_number(doc, "返回“训练历史”，确认刚才的记录状态为“已完成”，时间、案件和场景正确。")
    add_number(doc, "打开对话记录，核对学员消息、AI 回复、顺序和角色标识没有缺失或串到其他账号。")
    add_number(doc, "从历史记录再次打开评估报告，确认内容与刚生成时一致。")
    add_number(doc, "退出账号后重新登录，确认仍可找到本次记录。")
    add_callout(doc, "预期结果", "训练记录持久保存；对话与报告可再次访问；不同账号之间不可看到彼此的个人训练详情。", "success")


def add_special_steps(doc):
    add_heading(doc, "五、小组专项测试", 1)
    add_heading(doc, "5.1 B 组：语音输入与多角色", 2)
    add_bullet(doc, "切换到语音电话/语音输入，允许浏览器使用麦克风；完成至少 2 轮语音交流，再切回打字。")
    add_bullet(doc, "分别测试正常语速、1 次短暂停顿和1 次自我修正；不要故意大喊或播放无关音频。")
    add_bullet(doc, "观察识别文本是否可见、发送时机是否合理、能否关闭麦克风、语音结束后文本对话是否继续保持上下文。")
    add_bullet(doc, "多角色场景中，确认当前发言角色清楚，角色头像/标识变化合理，角色身份、立场和已知事实不串线。")
    add_callout(doc, "通过条件", "麦克风授权可控；语音识别可完成至少 2 轮；语音与文本可切换；AI 不因输入方式切换而丢失主要上下文。", "success")

    add_heading(doc, "5.2 C 组：中断与续训", 2)
    add_number(doc, "开始训练并完成 3 轮有效对话，记录最后一条消息。")
    add_number(doc, "通过页面正常返回或刷新页面模拟短暂中断；不要清除浏览器存储。")
    add_number(doc, "进入训练历史，找到“进行中”记录并点击“继续训练”。")
    add_number(doc, "核对中断前消息是否完整，继续发送 2 轮后完成训练并生成报告。")
    add_callout(doc, "通过条件", "进行中记录可找到；消息顺序和场景一致；不会生成多个重复活动会话；续训后仍能生成完整报告。", "success")

    add_heading(doc, "5.3 D 组：重复训练与评估可解释性", 2)
    add_number(doc, "完成第一轮后记录总评、3 个主要优点/不足和未命中考察点。")
    add_number(doc, "针对报告建议重新训练同一案件，主动补齐第一轮遗漏的关键问询或闭环动作。")
    add_number(doc, "比较两份报告：改进动作是否被识别；若分数变化，应能从对话和考察点找到原因。")
    add_callout(doc, "不是要求分数完全一致", "AI 评估允许存在合理波动，但同样的明显错误不应在一次被重罚、另一次完全忽略；明确补齐的关键动作不应毫无依据地导致更差评价。", "warning")

    add_heading(doc, "5.4 E 组：班级作业与视频实训（已配置时）", 2)
    add_bullet(doc, "班级作业：邀请码加入、查看公告与截止时间、开始/继续/完成作业，确认进度与完成数量同步更新。")
    add_bullet(doc, "视频实训：阅读入场说明，授权摄像头和麦克风，完成设备/人脸或在场校验后开始。")
    add_bullet(doc, "播放期间不拖动进度条；在训练节点按提示完成动作或话术，观察节点暂停、识别、反馈与继续播放。")
    add_bullet(doc, "完成后检查视频训练报告、历史记录与回放；若报告显示“待重修”，尝试重新训练入口。")
    add_callout(doc, "不适用处理", "若现场没有发布视频、浏览器策略禁止摄像头/麦克风，或组织者明确关闭视频实训，请标记“未配置/不适用”，不要计为功能失败。", "info")


def add_observation(doc):
    add_heading(doc, "七、测试时重点观察什么", 1)
    add_table(doc, ["维度", "重点问题", "良好表现"], [
        ["功能完整性", "能否完成登录、找任务、开始、对话、结束、报告、历史全链路？", "无阻断；状态前后一致；刷新/重登后数据仍在"],
        ["AI 角色真实感", "是否逐步释放信息，有情绪和立场，能持续 5 轮以上有效交流？", "回应针对当前问题；角色身份稳定；不机械重复"],
        ["专业与安全", "是否乱引法条、给出明显错误处置、虚构敏感事实或输出不当内容？", "不确定时谨慎表达；关键建议符合常识与规范；不泄露系统信息"],
        ["评估有效性", "评分、考察点和建议是否能对应具体对话？", "能指出完成/遗漏内容；建议具体；复训改进可被识别"],
        ["易用性", "首次使用者是否知道下一步做什么？按钮、状态和错误提示是否清楚？", "无需口头教学即可完成主流程；误操作可恢复"],
        ["性能稳定性", "多人同时提交时是否卡顿、超时、白屏、重复消息或报告失败？", "页面持续可用；失败有明确提示；无大面积 5xx/断线"],
        ["数据与权限", "是否看到他人的会话、报告或账号信息？输入的数据是否串号？", "只能访问本人内容；退出后受保护页面不可继续访问"],
        ["兼容性", "不同屏幕、Chrome/Edge、语音设备下是否可操作？", "核心按钮可见；中文不乱码；页面不重叠或横向溢出"],
    ], [1350, 4100, 3910], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.8)

    add_heading(doc, "不应直接判为系统缺陷的情况", 2)
    for text in [
        "组织者未发布案件、未给出班级邀请码或未配置视频内容。",
        "测试人员主动拒绝麦克风/摄像头权限，导致语音或视频识别不可用。",
        "测试人员网络断开、电脑休眠或浏览器被单位安全策略强制关闭。",
        "测试人员输入无意义字符、连续刷屏或在 AI 仍生成时反复提交。",
    ]:
        add_bullet(doc, text)


def add_acceptance(doc):
    add_heading(doc, "七、测试完成后要达到什么效果", 1)
    add_para(doc, "建议采用以下门槛判断本轮 50 人试测是否可以进入下一阶段。组织者应同时看“通过率”和“问题严重性”，不能只看平均分。")

    add_table(doc, ["验收项", "建议目标", "不通过信号"], [
        ["账号与登录", "≥49/50 可在 5 分钟内完成登录并进入正确学员端", "多人账号无效、角色跳错、反复登录或跨账号串号"],
        ["主流程完成率", "≥45/50 独立完成 1 次完整训练并生成报告", "5 人以上因同一系统问题无法完成"],
        ["对话持续性", "≥90% 的有效会话可持续 8 轮；至少 5 轮内无严重逻辑断裂", "普遍重复、答非所问、角色混淆或一次泄露全部案情"],
        ["响应体验", "普通页面约 3 秒内可用；AI 首次可见响应建议 ≤15 秒；报告建议 ≤60 秒", "集中超时、长时间无加载提示、重复响应或报告永久生成中"],
        ["评估质量", "≥80% 测试者认为报告与对话基本相关；主要扣分能找到依据", "大量模板化空话、分数与考察点矛盾、报告错配账号/会话"],
        ["历史与恢复", "完成记录 100% 可在历史中找到；C 组续训成功率 ≥9/10", "记录丢失、状态错误、继续训练创建错误会话"],
        ["并发稳定性", "50 人同时在线，≥40 人在 5 分钟内提交对话时，无全局不可用和数据错乱", "服务崩溃、大面积 5xx、消息串号或数据库明显异常"],
        ["严重缺陷", "P0=0；P1 在进入正式试点前清零或有经负责人批准的规避方案", "存在数据泄露、无法登录、主流程普遍阻断或结果错配"],
        ["易用性", "≥85% 测试者无需额外教学即可完成主流程；总体体验均分 ≥4.0/5", "多数人找不到下一步、关键按钮或报告入口"],
    ], [1700, 4080, 3580], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.7)

    add_heading(doc, "发布判断建议", 2)
    add_table(doc, ["结论", "判定建议"], [
        ["通过", "达到主要指标，P0/P1 为 0，可进入下一轮试点或扩大测试。"],
        ["有条件通过", "主流程可用，但存在少量 P1/P2；明确负责人、修复日期、规避方案和回归范围后继续。"],
        ["不通过", "出现 P0，或同一 P1 阻断 5 人以上，或主流程完成率低于 90%；修复后重新组织关键链路测试。"],
    ], [1700, 7660], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.2)


def add_defects(doc):
    add_heading(doc, "九、问题怎么记录与上报", 1)
    add_heading(doc, "8.1 严重程度", 2)
    add_table(doc, ["级别", "定义", "例子", "处理"], [
        ["P0 致命", "安全/隐私事故或平台整体不可用", "看到他人报告；50 人普遍无法登录；数据大面积错乱", "立即停止相关测试并通知负责人"],
        ["P1 严重", "主流程被阻断，且无可接受绕过方式", "无法开始/继续训练；报告持续失败；会话丢失", "现场立即报告，保留页面和时间"],
        ["P2 一般", "功能可继续，但结果不准确或体验明显受损", "筛选错误；语音偶发失败；评分依据不充分", "测试结束前提交完整记录"],
        ["P3 建议", "文案、视觉或低影响体验优化", "按钮名称不清楚；间距不一致；提示可更友好", "纳入体验改进清单"],
    ], [1050, 2650, 3540, 2120], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.6)

    add_heading(doc, "8.2 缺陷记录模板", 2)
    add_table(doc, ["字段", "填写内容"], [
        ["问题标题", "【级别】【模块】一句话描述，例如：【P1】【评估】报告生成超过 3 分钟后失败"],
        ["测试人员", "编号/小组；不要在公开群中填写身份证号等敏感信息"],
        ["发生时间", "精确到分钟，便于技术人员定位日志"],
        ["环境", "平台地址、浏览器及版本、电脑/手机、网络类型；语音/视频问题补充设备型号"],
        ["账号/会话", "测试账号名、案件/场景、训练记录编号（如页面可见）；不要填写密码"],
        ["前置条件", "例如：已登录、已完成 3 轮对话、使用语音模式"],
        ["复现步骤", "1……  2……  3……；写清点击了什么、输入了什么"],
        ["实际结果", "页面显示、错误提示、等待时间、是否可继续"],
        ["期望结果", "你认为系统应该怎样表现"],
        ["证据", "全屏截图/录屏、错误文字；截图前遮盖密码和不必要个人信息"],
        ["复现频率", "必现 / 3 次中 2 次 / 仅 1 次"],
    ], [1800, 7560], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.1)

    add_callout(doc, "遇到阻断时", "先记录时间和截图，再等待 30 秒并尝试一次正常返回/重新进入；仍失败则停止重复操作，联系现场技术保障。P0/P1 不要等到测试结束才上报。", "danger")


def add_questionnaire(doc):
    add_heading(doc, "十、测试结束反馈问卷", 1)
    add_para(doc, "评分说明：1=非常差/完全不同意，2=较差，3=一般，4=较好，5=非常好/完全同意。建议每名测试者在完成操作后立即填写。")
    add_table(doc, ["序号", "评价内容", "1", "2", "3", "4", "5"], [
        ["1", "我能在没有额外讲解的情况下完成管理员端建案和发布", "☐", "☐", "☐", "☐", "☐"],
        ["2", "案件、角色、场景、考察点和班级等管理功能容易理解", "☐", "☐", "☐", "☐", "☐"],
        ["3", "管理端配置后的内容能正确出现在配对学员端", "☐", "☐", "☐", "☐", "☐"],
        ["4", "训练页面结构、按钮和当前状态容易理解", "☐", "☐", "☐", "☐", "☐"],
        ["5", "AI 回答能针对我的问题，并保持角色和案情一致", "☐", "☐", "☐", "☐", "☐"],
        ["6", "系统响应速度可以接受，等待过程有明确反馈", "☐", "☐", "☐", "☐", "☐"],
        ["7", "评估报告能对应我的具体对话和处置动作", "☐", "☐", "☐", "☐", "☐"],
        ["8", "训练历史、对话回看和管理端状态便于我复盘/组织训练", "☐", "☐", "☐", "☐", "☐"],
        ["9", "两个端之间的数据、状态和权限边界符合我的预期", "☐", "☐", "☐", "☐", "☐"],
        ["10", "总体双端体验评分", "☐", "☐", "☐", "☐", "☐"],
    ], [600, 5760, 600, 600, 600, 600, 600], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 5, 8.8)

    add_heading(doc, "开放问题", 2)
    for prompt in [
        "最有帮助的一个功能是什么？为什么？",
        "最影响完成训练的一个问题是什么？",
        "AI 对话中最不像真人或最不专业的一处是什么？请尽量引用原话。",
        "评估报告中哪条结论最有用？哪条最不准确？",
        "如果只能改进一项，你最希望改进什么？",
    ]:
        add_para(doc, prompt, bold=True, color=NAVY, after=3)
        add_para(doc, "________________________________________________________________________________", color=MUTED, after=8)


def add_organizer_summary(doc):
    add_heading(doc, "十、组织者测试总结表", 1)
    add_table(doc, ["汇总项", "结果"], [
        ["实际参与人数", "________ 人；有效反馈 ________ 份"],
        ["成功登录人数", "________ / 50"],
        ["完整训练并生成报告人数", "________ / 50"],
        ["统一并发窗口", "开始 ________；结束 ________；峰值在线 ________；峰值并发对话 ________"],
        ["平均/典型 AI 响应时间", "________ 秒；异常最长 ________ 秒"],
        ["报告生成成功率", "________ %"],
        ["历史记录可回看率", "________ %"],
        ["总体体验平均分", "________ / 5"],
        ["缺陷数量", "P0 ____；P1 ____；P2 ____；P3 ____"],
        ["最终结论", "□ 通过    □ 有条件通过    □ 不通过"],
        ["下一步", "负责人：________  截止日期：________  回归范围：________________________"],
    ], [2500, 6860], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.3)

    add_heading(doc, "建议输出物", 2)
    for text in [
        "50 人参与与账号完成情况表。",
        "按 P0—P3 排序的缺陷清单，含负责人、修复版本和回归结果。",
        "AI 对话质量样本：优秀、一般、失败各 3—5 条脱敏会话。",
        "评估报告一致性对比与问卷统计。",
        "并发窗口的服务日志、接口错误率、响应时间和资源使用摘要。",
        "是否进入下一阶段的书面结论及遗留风险。",
    ]:
        add_bullet(doc, text)

    add_callout(doc, "完成定义", "只有当测试数据已汇总、P0/P1 已处理或有批准的规避方案、关键修复完成回归，并形成明确发布结论时，本轮 50 人测试才算真正结束。", "success")


def add_dual_end_scope(doc):
    add_heading(doc, "二、测试范围与执行原则", 1)
    add_callout(doc, "全员双端原则", "本轮不设专项小组。50 名测试人员均须完成管理端和学员端的完整测试闭环：先以管理员身份创建并发布自己的测试内容，再以学员身份完成训练、查看报告和复盘。", "success")
    add_table(doc, ["身份", "每人必须完成的核心事项", "完成凭证"], [
        ["管理员端", "登录、仪表盘、案件/场景/角色/考察点配置、发布、知识库、班级/学员、视频资源（若已启用）、个人中心", "测试编号对应的已发布案件/场景、截图、配置记录"],
        ["学员端", "登录、训练大厅/班级作业、文本对话、语音（已启用时）、训练结束、评估、历史/对话回看、设置", "本人的完整训练记录、评估报告、反馈问卷"],
        ["跨端闭环", "管理员发布的内容能被对应学员看到并完成；管理端统计或作业状态与学员端结果一致", "测试编号、案件名、场景名、训练记录号可对应"],
    ], [1300, 5360, 2700], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 9.1)
    add_heading(doc, "2.1 50 人账号与数据隔离规则", 2)
    add_table(doc, ["项目", "统一规则"], [
        ["账号配对", "每人领取一组账号：管理员 admin_test_01—admin_test_50；学员 student_test_01—student_test_50（实际命名可替换，但必须一一对应）。"],
        ["内容命名", "所有新建数据必须带个人编号：例如【T01】邻里纠纷、【T01】现场询问、【T01】报警人角色、【T01】班级。"],
        ["操作边界", "只能编辑、发布、删除本人编号的数据；不得修改其他编号的案件、角色、班级、学员或视频。"],
        ["测试材料", "组织者为每人准备同一套脱敏测试案情、考察点和知识文档，可复制使用但名称必须带个人编号。"],
        ["切换方式", "完成管理员端发布后先退出，再使用配对学员账号登录；不得用管理员账号替代学员端验收。"],
        ["环境要求", "必须在独立测试环境进行；本轮产生的 50 套数据在验收后按组织者方案统一归档或清理。"],
    ], [1700, 7660], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.1)
    add_heading(doc, "2.2 建议现场节奏（全员一致）", 2)
    add_table(doc, ["阶段", "建议时长", "全员统一动作"], [
        ["A. 入场与登录", "0—15 分钟", "领取双账号；登录管理端；确认仪表盘和权限。"],
        ["B. 管理端建案发布", "15—75 分钟", "创建案件、角色、场景和考察点；发布；完成知识库/班级等管理验证。"],
        ["C. 学员端训练", "75—120 分钟", "切换学员账号；找到本人发布的内容；完成至少 8 轮有效对话。"],
        ["D. 评估与复盘", "120—145 分钟", "生成报告；历史回看；退出重登；核对跨端状态。"],
        ["E. 并发与反馈", "145—165 分钟", "50 人同时在线提交对话或生成报告；填写问卷和缺陷单。"],
    ], [1700, 1600, 6060], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.0)


def add_admin_steps(doc):
    add_heading(doc, "四、怎么测试：管理端（全员必测）", 1)
    add_callout(doc, "目标", "验证教官/管理员能独立完成“准备内容—发布训练—组织学员—查看结果”的管理闭环；每一步都应使用本人测试编号，避免 50 人同时操作同一条数据。", "info")
    add_heading(doc, "步骤 1：管理员登录、仪表盘与权限", 2)
    for text in [
        "用本人管理员账号登录，确认自动进入管理端仪表盘，而不是学员端或维护端。",
        "查看仪表盘中的训练、学员或统计概览，刷新一次页面，确认无白屏、乱码和权限错误。",
        "尝试从浏览器直接访问维护端地址或维护端页面；管理员账号不应获得维护端账户管理权限。",
    ]:
        add_number(doc, text)
    add_callout(doc, "预期结果", "管理员登录成功、页面数据可加载；管理员仅进入管理端；受限页面得到明确的无权限提示或跳转，不泄露敏感数据。", "success")
    add_heading(doc, "步骤 2：案件、场景、角色与考察点", 2)
    add_table(doc, ["顺序", "操作", "重点验证与预期效果"], [
        ["2.1", "进入案件中心，新建【Txx】案件；录入名称、类型、背景、人员关系、冲突点、关键/隐藏事实等结构化内容。", "必填校验清楚；保存后重新打开信息不丢失；中文正常显示。"],
        ["2.2", "如提供 AI 解析/生成入口，使用组织者给出的脱敏材料尝试一次；检查生成内容并手工校对。", "生成过程有状态反馈；结果可编辑；不应覆盖已确认内容或生成明显无关信息。"],
        ["2.3", "为案件创建至少 2 个角色（如报警人、另一当事人），填写身份、诉求/顾虑、触发/安抚因素、开场和信息边界。", "角色可保存、编辑、绑定正确案件；角色字段不串到他人数据。"],
        ["2.4", "创建至少 1 个训练场景，设置名称、难度、描述、到场角色、主对话人和训练提示。", "场景角色校验有效；缺少主对话人或角色配置不合理时能提示。"],
        ["2.5", "创建或 AI 生成至少 3 个考察点；确认每项有名称、要求和可用于评估的内容。", "考察点可编辑、保存并在后续报告中出现。"],
        ["2.6", "发布本人场景；回到列表确认发布状态；如支持取消发布，记录入口但不要取消本次要验收的场景。", "已发布场景在学员端可见；未发布内容不应被学员端看到。"],
    ], [700, 4300, 4360], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.6)
    add_heading(doc, "步骤 3：知识库、学员、班级与作业", 2)
    for text in [
        "知识库：上传组织者提供的脱敏小文档，确认状态、文档信息和检索测试；使用与文档有关的问题，检查结果是否相关且无乱码。",
        "学员管理：搜索本人配对学员账号，确认基本信息与训练状态可查；不得修改或删除他人账号。若现场允许创建临时账号，名称必须使用【Txx】并记录后续清理。",
        "班级管理：创建【Txx】班级，将本人学员账号加入；发布一条简短公告；创建包含本人案件的作业并设置合理截止时间。",
        "回到班级/作业列表，确认班级成员、公告、案件、场景和完成规则保存正确。",
    ]:
        add_number(doc, text)
    add_callout(doc, "预期结果", "知识文档可上传和检索；班级邀请码/成员/作业可由对应学员使用；后台数据保存正确，不会误改其他测试人员的数据。", "success")
    add_heading(doc, "步骤 4：视频资源、数据质量与个人中心", 2)
    add_bullet(doc, "视频资源已启用时：进入视频库，查看已发布教学或互动视频的名称、封面、时长、场景和节点；按现场提供的测试资源创建/编辑或发布一项，确认学员端可见。")
    add_bullet(doc, "数据质量：打开数据质量页面，检查巡检结果、异常说明和筛选是否可用；只记录问题，不执行影响其他人训练数据的批量清理。")
    add_bullet(doc, "个人中心：查看个人信息、修改入口和密码策略提示；除非组织者要求，不要在测试中改动分配账号密码。")
    add_callout(doc, "未启用功能", "若视频实训或某项内容生产功能未在本轮环境配置，必须在反馈中标记“未配置/不适用”；只有已声明启用却无法使用时才记录为缺陷。", "warning")


def add_cross_terminal_steps(doc):
    add_heading(doc, "六、跨端闭环、并发与权限验证（全员必测）", 1)
    add_heading(doc, "6.1 管理端发布到学员端的闭环", 2)
    for text in [
        "退出管理员账号，以本人配对学员账号重新登录。",
        "在训练大厅搜索或筛选【Txx】案件；如创建了班级作业，使用邀请码加入【Txx】班级并从作业进入场景。",
        "确认案件标题、场景、难度、接警简报、角色和考察点所对应内容与管理端设置一致。",
        "完成训练、生成报告后重新以管理员账号登录，检查班级作业完成状态、学员训练状态或仪表盘统计是否同步。",
    ]:
        add_number(doc, text)
    add_callout(doc, "预期结果", "本人发布的场景仅在发布后可见；本人学员可以完成且记录归属正确；管理端看到的状态与学员端实际结果一致。", "success")
    add_heading(doc, "6.2 统一并发操作", 2)
    add_table(doc, ["时间点", "50 人统一动作", "现场记录"], [
        ["并发窗口开始", "50 人同时保持登录；确认管理端或学员端首页可操作。", "在线人数、加载失败人数、首屏耗时。"],
        ["第 1 分钟", "所有人进入本人已发布的训练场景。", "开始训练失败、重复会话、页面卡死。"],
        ["第 2—5 分钟", "所有人各发送 2 条有明确处置目标的文本消息；如已开启语音，使用语音者也统一发送。", "AI 首次响应、超时、重复回复、消息串号、错误码。"],
        ["第 6 分钟", "所有人点击结束训练/提交作业，集中生成评估报告。", "报告生成耗时、成功率、队列/错误提示。"],
        ["第 10 分钟", "所有人进入训练历史重新打开报告。", "记录与报告可访问率、数据一致性。"],
    ], [1400, 5120, 2840], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.8)
    add_heading(doc, "6.3 权限、数据隔离与异常恢复", 2)
    for text in [
        "使用学员账号尝试直接访问管理端页面；系统应拦截并返回学员工作区或登录页。",
        "使用管理员账号检查不能进入维护端账户管理；不要尝试绕过认证或猜测他人密码。",
        "在学员端完成 3 轮对话后正常刷新或返回历史，再继续训练；核对消息和进度是否保留。",
        "检查训练历史、对话记录和评估报告只能显示本人的内容；如发现可见他人内容，按 P0 立即上报。",
    ]:
        add_bullet(doc, text)


def add_dual_end_acceptance(doc):
    add_heading(doc, "八、测试完成后要达到什么效果", 1)
    add_para(doc, "本轮通过的核心不是只看“页面能打开”，而是 50 人均能用两种身份完成内容配置、发布、训练、评估和复盘，并且数据、权限与并发结果一致。")
    add_table(doc, ["验收项", "建议通过目标", "不通过信号"], [
        ["双账号登录", "≥49/50 测试人员可分别登录管理员端和学员端，且进入正确工作区", "角色跳错、多人无法登录、登录后反复跳转"],
        ["管理端建案发布", "≥45/50 独立创建、保存并发布 1 个带角色、场景和 ≥3 考察点的训练", "案件/场景保存失败；发布后学员端不可见；数据串号"],
        ["后台管理功能", "已启用的知识库、班级/作业、视频、数据质量、个人中心均可完成基本操作或正确显示", "已启用模块普遍白屏、权限错误、保存错误或明显数据丢失"],
        ["学员训练闭环", "≥45/50 用配对学员账号完成 ≥8 轮有效对话并成功生成报告", "训练/评估被同一问题阻断；报告无法生成或错配"],
        ["跨端一致性", "≥95% 配对数据可在两端正确对应；发布、作业状态、训练记录和报告一致", "未发布内容可见、已发布内容不可见、状态不同步、跨账号访问"],
        ["AI 与评估质量", "≥80% 测试者认为 AI 回应与报告基本贴合对话；考察点有完成/遗漏依据", "大量答非所问、角色混乱、模板空话、评分无依据"],
        ["并发稳定性", "50 人同时在线；≥40 人在 5 分钟内完成消息提交；报告窗口无全局不可用或数据错乱", "大面积 5xx、白屏、消息重复/串号、服务崩溃"],
        ["权限与隐私", "P0=0；学员不能进入管理端，管理员不能进入维护端；任何人不能读取他人训练详情", "越权可见、数据泄露、跨账号修改、密码/令牌泄露"],
        ["易用性", "≥85% 测试者无需额外教学完成两端关键路径；总体评分 ≥4.0/5", "多数人不清楚发布、开始训练、报告或历史入口"],
    ], [1650, 4250, 3460], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], 8.5)
    add_callout(doc, "结论规则", "出现任一 P0（越权、数据泄露、数据串号、整体不可用）即本轮不通过；同一 P1 导致 5 人以上无法完成管理端发布或学员端训练，也应修复后重新测试关键闭环。", "danger")


def add_dual_end_summary(doc):
    add_heading(doc, "十一、组织者双端测试总结表", 1)
    add_table(doc, ["汇总项", "结果"], [
        ["实际测试人数", "________ / 50；管理员端完成 ________；学员端完成 ________"],
        ["双账号登录成功", "管理员 ________ / 50；学员 ________ / 50"],
        ["案件/场景发布成功", "________ / 50；学员端可见 ________ / 50"],
        ["知识库/班级/视频等已启用模块", "模块：________；基本操作成功率：________ %"],
        ["完整训练与报告", "________ / 50；历史可回看 ________ / 50"],
        ["跨端数据一致性", "________ / 50；不一致说明：________________________"],
        ["并发窗口", "峰值在线 ________；消息成功 ________ / 50；报告成功 ________ / 50"],
        ["体验平均分", "管理端 ________ / 5；学员端 ________ / 5"],
        ["缺陷数量", "P0 ____；P1 ____；P2 ____；P3 ____"],
        ["最终结论", "□ 通过    □ 有条件通过    □ 不通过"],
        ["下一步", "负责人：________  修复截止：________  回归范围：________________________"],
    ], [2600, 6760], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], 9.1)
    add_callout(doc, "完成定义", "只有在 50 人双端数据已汇总、P0/P1 已处理或有批准规避方案、关键修复已回归、并形成清晰发布结论后，本轮双端全员测试才算完成。", "success")


def build():
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "AI虚拟警情处置模拟训练平台 50人测试使用说明"
    props.subject = "50人集中测试执行、验收与反馈指南"
    props.author = "项目测试组"
    props.keywords = "警情处置, 模拟训练, 用户测试, 验收, 50人"

    add_cover(doc)
    add_quick_card(doc)
    add_dual_end_scope(doc)
    add_preparation(doc)
    add_admin_steps(doc)
    add_steps(doc)
    add_cross_terminal_steps(doc)
    add_observation(doc)
    add_dual_end_acceptance(doc)
    add_defects(doc)
    add_questionnaire(doc)
    add_dual_end_summary(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
